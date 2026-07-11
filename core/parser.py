from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZipFile
import os
import subprocess
import tempfile
import xml.etree.ElementTree as ET

import docx


WORD_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
PARA_TAG = WORD_NS + "p"
TEXT_TAG = WORD_NS + "t"


def extract_resume_text(path: str | Path) -> str:
    resume_path = Path(path)
    suffix = resume_path.suffix.lower()

    if suffix == ".docx":
        return clean_text(extract_docx_text(resume_path))
    if suffix == ".pdf":
        return clean_text(extract_pdf_text(resume_path))
    if suffix in {".txt", ".md"}:
        return clean_text(resume_path.read_text(encoding="utf-8", errors="ignore"))

    raise ValueError(f"Unsupported resume format: {suffix}. Use DOCX, PDF, TXT, or MD.")


def extract_docx_text(path: str | Path) -> str:
    return _extract_docx_content_text(Path(path).read_bytes())


def extract_pdf_text(path: str | Path) -> str:
    from pypdf import PdfReader

    text_parts = []
    resume_path = Path(path)
    reader = PdfReader(str(resume_path))
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    extracted = clean_text("\n\n".join(text_parts))
    return extracted or _extract_pdf_text_with_optional_ocr(resume_path)


def extract_uploaded_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return clean_text(_extract_docx_content_text(content))
    if suffix in {".txt", ".md"}:
        return clean_text(content.decode("utf-8", errors="ignore"))
    if suffix == ".pdf":
        from pypdf import PdfReader

        text_parts = []
        reader = PdfReader(BytesIO(content))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        extracted = clean_text("\n\n".join(text_parts))
        if extracted:
            return extracted
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(content)
            tmp_path = Path(tmp.name)
        try:
            return clean_text(_extract_pdf_text_with_optional_ocr(tmp_path))
        finally:
            try:
                tmp_path.unlink()
            except OSError:
                pass
    raise ValueError(f"Unsupported upload format: {suffix}. Use DOCX, PDF, TXT, or MD.")


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    cleaned_lines: list[str] = []
    prev_empty = False
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            if not prev_empty:
                cleaned_lines.append("")
            prev_empty = True
        else:
            cleaned_lines.append(stripped)
            prev_empty = False
    return "\n".join(cleaned_lines).strip()


def _extract_docx_content_text(content: bytes) -> str:
    structured_text = ""
    try:
        document = docx.Document(BytesIO(content))
        structured_text = _extract_document_text(document)
    except Exception:
        structured_text = ""
    xml_text = _extract_docx_xml_text(content)
    return _merge_text_sources(structured_text, xml_text)


def _extract_document_text(document: docx.Document) -> str:
    parts: list[str] = []
    parts.extend(p.text for p in document.paragraphs if p.text.strip())
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    for section in document.sections:
        parts.extend(p.text for p in section.header.paragraphs if p.text.strip())
        parts.extend(p.text for p in section.footer.paragraphs if p.text.strip())
    return "\n".join(parts)


def _extract_docx_xml_text(content: bytes) -> str:
    parts: list[str] = []
    try:
        with ZipFile(BytesIO(content)) as archive:
            xml_names = [
                name
                for name in archive.namelist()
                if name.startswith("word/")
                and name.endswith(".xml")
                and not name.startswith("word/styles")
                and not name.startswith("word/settings")
                and not name.startswith("word/fontTable")
            ]
            for name in xml_names:
                try:
                    root = ET.fromstring(archive.read(name))
                except ET.ParseError:
                    continue
                for paragraph in root.iter(PARA_TAG):
                    texts = [node.text or "" for node in paragraph.iter(TEXT_TAG)]
                    line = "".join(texts).strip()
                    if line:
                        parts.append(line)
    except Exception:
        return ""
    return "\n".join(parts)


def _merge_text_sources(*sources: str) -> str:
    seen: set[str] = set()
    merged: list[str] = []
    for source in sources:
        for raw_line in source.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            key = " ".join(line.split())
            if key in seen:
                continue
            seen.add(key)
            merged.append(line)
    return "\n".join(merged)


def _extract_pdf_text_with_optional_ocr(path: Path) -> str:
    command = os.environ.get("GENERIC_REPORT_OCR_COMMAND", "").strip()
    if not command:
        return ""
    args = [part.format(input=str(path)) for part in command.split()]
    if "{input}" not in command:
        args.append(str(path))
    try:
        completed = subprocess.run(
            args,
            text=True,
            encoding="utf-8",
            errors="replace",
            capture_output=True,
            timeout=int(os.environ.get("GENERIC_REPORT_OCR_TIMEOUT", "120")),
            check=False,
        )
    except Exception:
        return ""
    if completed.returncode != 0:
        return ""
    return completed.stdout or ""


__all__ = [
    "extract_resume_text",
    "extract_docx_text",
    "extract_pdf_text",
    "extract_uploaded_text",
    "clean_text",
]
