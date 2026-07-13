"""
DOCX 渲染引擎 - 支持"填充现有模板"和"程序化构建"两种模式
"""

from __future__ import annotations

import base64
import re
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from .placeholder_report import build_placeholder_context
from .resume_parser import resume_sections_from_data


TSTAR_CN = "\u6cf0\u4f26\u4ed5"


class ReportRenderer:
    """报告渲染引擎"""

    def __init__(self, brand_config: dict[str, Any], template_config: dict[str, Any] | None = None):
        self.brand_config = brand_config
        self.template_config = template_config or {}
        self.render_rules = self.template_config.get("render_rules", {})
        # 从 brand_config 的 branding 中读取字体配置作为 fallback
        branding = brand_config.get("branding", {})
        self.font_config = self.render_rules.get("font", {}) or {
            "family": branding.get("font_family", "Microsoft YaHei"),
            "family_en": branding.get("font_family_en", "Arial"),
            "size": branding.get("font_size", 10.5),
        }

    def render(self, data: dict[str, Any], output_path: str | Path) -> Path:
        """
        渲染报告

        Args:
            data: 报告数据
            output_path: 输出文件路径

        Returns:
            输出文件路径
        """
        output_path = Path(output_path)

        # 判断使用哪种模式
        template_mapping = self.brand_config.get("template_mapping", {})
        use_client_template = template_mapping.get("use_client_template", False)

        if use_client_template and self.template_config:
            # 模式 1: 填充客户提供的模板
            doc = self._render_with_template(data)
        else:
            # 模式 2: 程序化构建
            doc = self._render_programmatically(data)

        # 保存
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        return output_path

    # ============================================================
    # 模式 1: 填充现有模板
    # ============================================================

    def _render_with_template(self, data: dict[str, Any]) -> Document:
        """使用客户提供的模板填充数据"""
        template_mapping = self.brand_config.get("template_mapping", {})
        template_path = template_mapping.get("client_template_path", "")

        if not template_path:
            raise ValueError("client_template_path is not configured")

        # 加载模板
        full_path = Path(__file__).parent.parent / "templates" / template_path
        if not full_path.exists():
            raise FileNotFoundError(f"Template file not found: {full_path}")

        doc = Document(str(full_path))

        # 处理占位符替换
        self._replace_placeholders(doc, data)

        # 处理字段映射
        field_mappings = self.template_config.get("field_mappings", [])
        for mapping in field_mappings:
            self._apply_field_mapping(doc, mapping, data)

        # 处理条件渲染
        conditional_rules = self.template_config.get("conditional_rendering", [])
        for rule in conditional_rules:
            self._apply_conditional_rendering(doc, rule, data)

        # 应用全局样式
        self._apply_global_styles(doc)

        return doc

    def _replace_placeholders(self, doc: Document, data: dict[str, Any]) -> None:
        """替换 {{placeholder}} 形式的占位符"""
        placeholders = self.template_config.get("placeholders", [])

        # 构建占位符映射
        placeholder_map = {}
        for ph in placeholders:
            placeholder = ph.get("placeholder", "")
            field = ph.get("field", "")
            if not placeholder or not field:
                continue

            # 获取值
            value = self._get_nested_value(data, field)
            if value is not None:
                placeholder_map[placeholder] = self._format_value(value, ph.get("format"))

        # 替换段落中的占位符
        for paragraph in doc.paragraphs:
            for placeholder, value in placeholder_map.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in placeholder_map.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))

    def _apply_field_mapping(self, doc: Document, mapping: dict[str, Any], data: dict[str, Any]) -> None:
        """应用字段映射到文档"""
        field_name = mapping.get("field")
        target = mapping.get("target")
        render_mode = mapping.get("render_mode", "text")

        if not field_name or not target:
            return

        # 获取字段值
        field_value = data.get(field_name)
        if field_value is None:
            return

        # 解析目标位置
        target_parts = target.split(".")

        if target_parts[0] == "table":
            self._fill_table_cell(doc, target_parts, field_value, render_mode, mapping)
        elif target_parts[0] == "paragraph":
            self._fill_paragraph(doc, target_parts, field_value)

    def _fill_table_cell(
        self,
        doc: Document,
        target_parts: list[str],
        value: Any,
        render_mode: str,
        mapping: dict[str, Any],
    ) -> None:
        """填充表格单元格"""
        try:
            table_idx = int(target_parts[1])
            if table_idx >= len(doc.tables):
                return

            table = doc.tables[table_idx]

            # 如果目标是整个表格（如原始简历）
            if len(target_parts) == 2:
                if render_mode == "append_paragraphs":
                    self._append_to_table(table, value, mapping)
                return

            row_idx = int(target_parts[3]) if len(target_parts) > 3 else 0
            cell_idx = int(target_parts[5]) if len(target_parts) > 5 else 0

            if row_idx >= len(table.rows):
                return

            cell = table.rows[row_idx].cells[cell_idx]

            if render_mode == "structured":
                self._fill_structured_cell(cell, value, mapping.get("structured_format", {}))
            else:
                self._set_cell_text(cell, str(value))

        except (IndexError, ValueError) as e:
            print(f"Warning: failed to fill table cell at {'.'.join(target_parts)}: {e}")

    def _fill_structured_cell(self, cell: Any, value: dict[str, Any], format_config: dict[str, Any]) -> None:
        """填充结构化字段（如 Recommendation Rationale）"""
        # 清空单元格
        cell.text = ""

        header = format_config.get("header", "")
        sub_fields = format_config.get("sub_fields", [])
        separator = format_config.get("separator", "\n\n")
        sub_field_separator = format_config.get("sub_field_separator", "\n")

        paragraphs = []

        if header:
            paragraphs.append(f"{header}")
            paragraphs.append("")

        for sub_def in sub_fields:
            sub_field_name = sub_def.get("field")
            sub_label = sub_def.get("label", sub_field_name)
            prefix = sub_def.get("prefix", "")
            condition = sub_def.get("condition")

            sub_value = value.get(sub_field_name) if isinstance(value, dict) else None

            # 条件检查
            if condition == "not_empty" and not sub_value:
                continue

            if sub_value:
                formatted = f"{prefix}{sub_label}: {sub_value}"
                paragraphs.append(formatted)

        # 写入单元格
        cell.text = "\n".join(paragraphs)
        self._apply_font_to_cell(cell)

    def _append_to_table(self, table: Any, value: str, mapping: dict[str, Any]) -> None:
        """将长文本分段追加到表格"""
        config = mapping.get("append_paragraphs_config", {})
        max_chars = config.get("max_chars_per_cell", 2000)
        split_strategy = config.get("split_strategy", "paragraph")

        if split_strategy == "paragraph":
            paragraphs = value.split("\n\n")
        else:
            paragraphs = [value[i:i + max_chars] for i in range(0, len(value), max_chars)]

        for para in paragraphs:
            if not para.strip():
                continue

            row = table.add_row()
            cell = row.cells[0]
            cell.text = para.strip()
            self._apply_font_to_cell(cell)

    def _fill_paragraph(self, doc: Document, target_parts: list[str], value: Any) -> None:
        """填充段落"""
        try:
            para_idx = int(target_parts[1])
            if para_idx < len(doc.paragraphs):
                doc.paragraphs[para_idx].text = str(value)
        except (IndexError, ValueError):
            pass

    def _apply_conditional_rendering(self, doc: Document, rule: dict[str, Any], data: dict[str, Any]) -> None:
        """应用条件渲染规则"""
        condition = rule.get("condition")
        field_name = rule.get("field")
        action = rule.get("action")
        target = rule.get("target")

        if not all([condition, field_name, action, target]):
            return

        # 获取字段值
        field_value = data.get(field_name)

        # 评估条件
        condition_met = False
        if condition == "field_empty":
            condition_met = not field_value or (isinstance(field_value, str) and not field_value.strip())
        elif condition == "field_not_empty":
            condition_met = bool(field_value) and (not isinstance(field_value, str) or field_value.strip())

        if not condition_met:
            return

        # 执行动作
        if action == "hide_row":
            self._hide_row(doc, target)
        elif action == "show_placeholder":
            self._show_placeholder(doc, target, rule.get("placeholder", ""))

    def _hide_row(self, doc: Document, target: str) -> None:
        """隐藏表格行"""
        try:
            parts = target.split(".")
            table_idx = int(parts[1])
            row_idx = int(parts[3])

            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                if row_idx < len(table.rows):
                    row = table.rows[row_idx]
                    # 隐藏行：将行高设为 0
                    row.height = Pt(0)
                    # 清空内容
                    for cell in row.cells:
                        cell.text = ""
        except (IndexError, ValueError):
            pass

    def _show_placeholder(self, doc: Document, target: str, placeholder: str) -> None:
        """显示占位符"""
        try:
            parts = target.split(".")
            table_idx = int(parts[1])
            row_idx = int(parts[3])
            cell_idx = int(parts[5]) if len(parts) > 5 else 0

            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                if row_idx < len(table.rows):
                    cell = table.rows[row_idx].cells[cell_idx]
                    cell.text = placeholder
                    self._apply_font_to_cell(cell, italic=True, color="999999")
        except (IndexError, ValueError):
            pass

    # ============================================================
    # 模式 2: 程序化构建
    # ============================================================

    def _render_programmatically(self, data: dict[str, Any]) -> Document:
        """程序化构建 DOCX"""
        if self.brand_config.get("brand_id") == "tstar":
            return self._render_tstar_report_v2(data)

        return self._render_default_report(data)

    def _render_tstar_report_v2(self, data: dict[str, Any]) -> Document:
        """Render the T-STAR report using the approved burgundy business template."""
        ctx = build_placeholder_context(data, self.brand_config)
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        self._apply_document_font(doc)
        self._doc_appendix_footer(doc)

        primary = "8F0E5C"
        accent = "8F0E5C"
        border = "E7D3DF"
        logo_border = "E4C567"

        header = doc.add_table(rows=1, cols=2)
        header.alignment = WD_TABLE_ALIGNMENT.CENTER
        header.autofit = False
        header.columns[0].width = Cm(4.5)
        header.columns[1].width = Cm(12.5)
        self._remove_table_borders(header)

        logo_cell = header.rows[0].cells[0]
        info_cell = header.rows[0].cells[1]
        logo_cell.text = ""
        info_cell.text = ""
        self._set_cell_border_v2(logo_cell, logo_border)
        self._set_cell_vertical_alignment(logo_cell)
        self._add_logo_to_paragraph(logo_cell.paragraphs[0], width=1.45)
        cn = logo_cell.add_paragraph()
        cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_run(cn, TSTAR_CN, size=8, bold=True, color="111111")

        title = info_cell.paragraphs[0]
        title.paragraph_format.space_before = Pt(0)
        title.paragraph_format.space_after = Pt(5)
        self._add_run(title, "Candidate Referral Report / 候选人推荐报告", size=18, bold=True, color="111827")

        meta_rows = 3 if ctx.get("salary_info") else 2
        meta_table = info_cell.add_table(rows=meta_rows, cols=2)
        meta_table.autofit = False
        meta_table.columns[0].width = Cm(6.2)
        meta_table.columns[1].width = Cm(6.4)
        self._remove_table_borders(meta_table)
        self._set_meta_cell(meta_table.rows[0].cells[0], "Candidate", ctx["candidate_name"], accent)
        self._set_meta_cell(meta_table.rows[0].cells[1], "Target Role", ctx["target_role"], accent)
        self._set_meta_cell(meta_table.rows[1].cells[0], "Client", ctx["client_company"], accent)
        self._set_meta_cell(meta_table.rows[1].cells[1], "Current", ctx["current_summary"], accent)
        if ctx.get("salary_info"):
            self._set_meta_cell(meta_table.rows[2].cells[0], "Salary", ctx["salary_info"], accent)
            meta_table.rows[2].cells[1].text = ""

        rule = doc.add_paragraph()
        rule.paragraph_format.space_before = Pt(6)
        rule.paragraph_format.space_after = Pt(8)
        self._doc_rule(rule, primary)

        self._doc_heading(doc, "Candidate Profile / 候选人基本信息", accent)
        if ctx.get("personal_info_rows"):
            personal_rows = list(ctx["personal_info_rows"])
            self._doc_personal_info(
                doc,
                personal_rows,
                border,
                accent,
                ctx.get("professional_photo_data_uri", ""),
                bool(ctx.get("professional_photo_required")),
            )
        else:
            self._doc_panel_v2(doc, [("Name / 姓名", ctx["candidate_name"]), ("Current / 当前", ctx["current_summary"])], border)

        self._doc_heading(doc, "Recommendation Summary / 推荐摘要", accent)
        self._doc_panel_v2(
            doc,
            [
                ("Motivation / 求职动机", ctx["motivation"]),
                ("Role Fit / 岗位匹配", ctx["role_fit"]),
            ],
            border,
        )

        self._doc_heading(doc, "Consultant Assessment / 顾问评估", accent)
        self._doc_panel_v2(
            doc,
            [
                ("Strengths / 推荐亮点", ctx["strengths_summary"]),
                ("Risks / Questions / 风险与待确认", ctx["risk_notes"]),
            ],
            border,
        )

        if ctx["appendix_blocks"].get("experience_groups"):
            self._doc_heading(doc, "Work Experience / 工作经历", accent)
            self._doc_experience_groups(doc, ctx["appendix_blocks"]["experience_groups"], accent)

        if ctx["job_description"]:
            self._doc_heading(doc, "Role Requirement Notes / JD 要求", accent)
            self._text_box_v2(doc, str(ctx["job_description"]), border)

        doc.add_page_break()
        self._doc_heading(doc, "Original Resume Appendix / 原始简历附录", accent)
        self._original_resume_appendix(doc, ctx["appendix_resume"])
        return doc

    def _render_tstar_report(self, data: dict[str, Any]) -> Document:
        """Render a T-STAR bilingual referral report with the approved template structure."""
        ctx = build_placeholder_context(data, self.brand_config)
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        self._apply_document_font(doc)
        self._doc_appendix_footer(doc)

        primary = "D6A100"
        accent = "8A6A00"
        border = "E8D9A8"

        cover_title = doc.add_paragraph()
        cover_title.paragraph_format.space_after = Pt(2)
        self._add_run(cover_title, "T-STAR 泰伦仕 候选人推荐报告双语模板", size=20, bold=True, color=primary)
        cover_subtitle = doc.add_paragraph()
        cover_subtitle.paragraph_format.space_after = Pt(12)
        self._add_run(cover_subtitle, "Candidate Recommendation Report Template", size=10.5, color="374151")
        self._doc_rule(cover_title, primary)
        self._doc_logo_card(doc, primary)
        self._doc_heading(doc, "1. 封面与使用说明", primary)
        self._doc_usage_columns(doc, primary, border)

        doc.add_page_break()

        header = doc.add_table(rows=1, cols=2)
        header.alignment = WD_TABLE_ALIGNMENT.CENTER
        header.autofit = False
        header.columns[0].width = Cm(4.2)
        header.columns[1].width = Cm(12.8)
        self._remove_table_borders(header)

        logo_cell = header.rows[0].cells[0]
        info_cell = header.rows[0].cells[1]
        logo_cell.text = ""
        info_cell.text = ""
        self._set_cell_border_v2(logo_cell, border)
        self._set_cell_vertical_alignment(logo_cell)
        self._add_logo_to_paragraph(logo_cell.paragraphs[0], width=1.45)
        cn = logo_cell.add_paragraph()
        cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_run(cn, TSTAR_CN, size=8, bold=True, color="111111")

        title = info_cell.paragraphs[0]
        title.paragraph_format.space_before = Pt(0)
        title.paragraph_format.space_after = Pt(5)
        self._add_run(title, "Candidate Referral Report / 候选人推荐报告", size=18, bold=True, color="111111")

        meta_rows = 3 if ctx.get("salary_info") else 2
        meta_table = info_cell.add_table(rows=meta_rows, cols=2)
        meta_table.autofit = False
        meta_table.columns[0].width = Cm(6.2)
        meta_table.columns[1].width = Cm(6.4)
        self._remove_table_borders(meta_table)
        self._set_meta_cell(meta_table.rows[0].cells[0], "Candidate", ctx["candidate_name"], accent)
        self._set_meta_cell(meta_table.rows[0].cells[1], "Target Role", ctx["target_role"], accent)
        self._set_meta_cell(meta_table.rows[1].cells[0], "Client", ctx["client_company"], accent)
        self._set_meta_cell(meta_table.rows[1].cells[1], "Current", ctx["current_summary"], accent)
        if ctx.get("salary_info"):
            self._set_meta_cell(meta_table.rows[2].cells[0], "Salary", ctx["salary_info"], accent)
            meta_table.rows[2].cells[1].text = ""

        rule = doc.add_paragraph()
        rule.paragraph_format.space_before = Pt(6)
        rule.paragraph_format.space_after = Pt(8)
        self._doc_rule(rule, primary)

        self._doc_heading(doc, "Profile Summary / 候选人概要", accent)
        self._doc_panel_v2(
            doc,
            [
                ("Motivation / 求职动机", ctx["motivation"]),
                ("Role Fit / 岗位匹配", ctx["role_fit"]),
            ],
            border,
        )

        self._doc_heading(doc, "Consultant Assessment / 顾问评估", accent)
        self._doc_panel_v2(
            doc,
            [
                ("Strengths / 推荐亮点", ctx["strengths_summary"]),
                ("Risks / Questions / 风险与待确认", ctx["risk_notes"]),
            ],
            border,
        )

        if ctx["appendix_blocks"].get("experience_groups"):
            self._doc_heading(doc, "Work Experience / 工作经历", accent)
            self._doc_experience_groups(doc, ctx["appendix_blocks"]["experience_groups"], accent)

        if ctx["job_description"]:
            self._doc_heading(doc, "Role Requirement Notes / JD 要求", accent)
            self._text_box_v2(doc, str(ctx["job_description"]), border)

        doc.add_page_break()
        self._doc_heading(doc, "Original Resume Appendix / 原始简历附录", accent)
        self._doc_paragraph(doc, ctx["appendix_resume"] or ctx["original_resume"] or "Original resume not parsed.", size=9.5)
        return doc

    def _render_default_report(self, data: dict[str, Any]) -> Document:
        """默认品牌报告渲染"""
        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.font_config.get("family_en", "Arial")
        font.size = Pt(self.font_config.get("size", 10.5))

        # 添加标题
        title = doc.add_heading("候选人推荐报告", level=0)
        title.alignment = 1  # 居中

        # 添加基本信息
        doc.add_heading("职位信息", level=1)
        self._add_info_table(doc, data, ["position_title", "location", "department_function", "req_id"])

        # 添加候选人信息
        doc.add_heading("候选人信息", level=1)
        self._add_info_table(doc, data, ["candidate_name", "current_title", "current_company", "years_of_experience"])

        # 添加顾问评估
        doc.add_heading("顾问评估", level=1)
        self._add_assessment_section(doc, data)

        # 添加原始简历
        if data.get("original_resume"):
            doc.add_heading("原始简历", level=1)
            doc.add_paragraph(data["original_resume"])

        return doc

    # ============================================================
    # T-STAR 辅助方法
    # ============================================================

    def _tstar_logo_path(self) -> Path:
        logo_dir = Path(__file__).parent.parent / "templates" / "tstar"
        for name in (
            "tstar_logo_white.png",
            "tstar_logo_white.jpg",
            "tstar_logo.png",
            "tstar_logo.jpg",
        ):
            path = logo_dir / name
            if path.exists():
                return path
        return logo_dir / "tstar_logo.jpg"

    def _add_logo_to_paragraph(self, paragraph: Any, width: float = 2.5) -> None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_path = self._tstar_logo_path()
        if logo_path.exists():
            paragraph.add_run().add_picture(str(logo_path), width=Inches(width))
            return
        self._add_run(paragraph, "T-STAR 泰伦仕", size=14, bold=True, color="111111")

    def _doc_rule(self, paragraph: Any, color: str) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "18")
        bottom.set(qn("w:space"), "6")
        bottom.set(qn("w:color"), color)
        border.append(bottom)
        p_pr.append(border)

    def _doc_logo_card(self, doc: Document, color: str) -> None:
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Cm(12.5)
        cell = table.rows[0].cells[0]
        cell.text = ""
        self._set_cell_border_v2(cell, "E5E7EB")
        p = cell.paragraphs[0]
        self._add_run(p, "T-STAR 官方标准推荐", size=10, bold=True, color="111827")
        logo_p = cell.add_paragraph()
        self._add_logo_to_paragraph(logo_p, width=2.7)
        caption = cell.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_run(caption, "专注高端人才寻访与专业招聘咨询", size=9.5, color="374151")

    def _doc_usage_columns(self, doc: Document, color: str, border: str) -> None:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(8.2)
        table.columns[1].width = Cm(8.2)
        self._fill_usage_cell(
            table.rows[0].cells[0],
            "中文说明",
            [
                ("模板定位", "本模板用于 T-STAR 泰伦仕候选人推荐报告交付，采用“主报告 + 原始简历附录 + 解析置信度/人工复核规则”的结构。主报告聚焦客户决策所需的核心信息，附录保留原始简历证据。"),
                ("使用原则", "顾问应优先填写高置信度字段，并基于面试、JD 和原始简历证据完善推荐理由。对低置信度内容应明确标记为待复核，避免过度改写原始材料。"),
                ("交付要求", "导出前请完成占位符检查、候选人信息核对、客户公司/目标岗位确认，并确保附录或原始文件附件已随报告一并交付。"),
            ],
            color,
            border,
        )
        self._fill_usage_cell(
            table.rows[0].cells[1],
            "English Notes",
            [
                ("Template Positioning", "This template supports T-STAR candidate recommendation delivery through a main report, an original resume appendix, and parsing confidence / manual review rules. The main report highlights decision-ready evidence while the appendix preserves source material."),
                ("Usage Principles", "Consultants should complete high-confidence fields first and refine recommendation rationale with interview notes, JD requirements, and resume evidence. Low-confidence content should be clearly marked for review instead of over-rewritten."),
                ("Delivery Checklist", "Before release, check all placeholders, validate candidate details, confirm client company and target role, and ensure the appendix or original source attachment is included with the report."),
            ],
            color,
            border,
        )

    def _fill_usage_cell(self, cell: Any, heading: str, notes: list[tuple[str, str]], color: str, border: str) -> None:
        cell.text = ""
        self._set_cell_border_v2(cell, border)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = cell.paragraphs[0]
        self._add_run(p, heading, size=12, bold=True, color=color)
        for label, body in notes:
            hp = cell.add_paragraph()
            hp.paragraph_format.space_before = Pt(5)
            hp.paragraph_format.space_after = Pt(1)
            self._add_run(hp, label, size=9.5, bold=True, color="111827")
            bp = cell.add_paragraph()
            bp.paragraph_format.space_after = Pt(4)
            self._add_run(bp, body, size=8.8, color="374151")

    def _doc_appendix_footer(self, doc: Document) -> None:
        text = (
            "附录占位提示：此处插入原始简历全文或原始文件附件。 | "
            "Appendix Placeholder: Insert original resume content or attached source file here."
        )
        text = (
            "附录提示：原始简历全文仅保留在附录或随报告附件交付。 | "
            "Appendix: the original resume is preserved only in the appendix or as an attached source file."
        )
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.text = ""
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_run(paragraph, text, size=8, color="6B7280")

    def _doc_confidence_rules(self, doc: Document, color: str, border: str) -> None:
        self._doc_heading(doc, "Parsing Confidence & Manual Review Rules / 解析置信度与人工复核规则", color)
        rows = [
            ("High Confidence / 高置信", "Names, contact details, education, company names, title names, and explicit date ranges may be extracted automatically with consultant spot-check."),
            ("Medium Confidence / 中置信", "Skills, certificates, reporting lines, language level, and project periods require consultant validation against interview notes and source materials."),
            ("Low Confidence / 低置信", "Achievements, role motivation, compensation, departure reason, and nuanced leadership impact must be reviewed or rewritten by consultants before client submission."),
        ]
        table = doc.add_table(rows=0, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(5.2)
        table.columns[1].width = Cm(11.8)
        for label, body in rows:
            row = table.add_row()
            for cell in row.cells:
                self._set_cell_border_v2(cell, border)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            row.cells[0].text = ""
            row.cells[1].text = ""
            self._add_run(row.cells[0].paragraphs[0], label, size=9.2, bold=True, color=color)
            self._add_run(row.cells[1].paragraphs[0], body, size=9.2, color="374151")

    def _set_meta_cell(self, cell: Any, label: str, value: str, color: str) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        self._add_run(p, f"{label}: ", size=9, bold=True, color=color)
        self._add_run(p, value or "-", size=9, color="374151")

    def _doc_heading(self, doc: Document, text: str, color: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        self._add_run(p, text, size=12, bold=True, color=color)

    def _doc_panel_v2(self, doc: Document, rows: list[tuple[str, str]], border: str) -> None:
        """Improved panel with lighter borders and better spacing."""
        table = doc.add_table(rows=0, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(2.5)
        table.columns[1].width = Cm(14.5)
        for label, value in rows:
            row = table.add_row()
            label_cell, value_cell = row.cells
            for item in row.cells:
                self._set_cell_border_v2(item, border)
                self._set_cell_vertical_alignment(item)
            label_cell.text = ""
            value_cell.text = ""
            self._add_run(label_cell.paragraphs[0], label, size=9, bold=True, color="8F0E5C")
            self._add_run(value_cell.paragraphs[0], value or "-", size=9.5, color="374151")

    def _doc_personal_info_legacy_unused(
        self,
        doc: Document,
        rows: list[tuple[str, str]],
        border: str,
        color: str,
        photo_data_uri: str = "",
        photo_required: bool = False,
    ) -> None:
        show_photo_area = bool(str(photo_data_uri or "").strip()) or bool(photo_required)
        layout = doc.add_table(rows=1, cols=2 if show_photo_area else 1)
        layout.autofit = False
        if show_photo_area:
            layout.columns[0].width = Cm(3.4)
            layout.columns[1].width = Cm(13.6)
            photo_cell, info_cell = layout.rows[0].cells
            cells = (photo_cell, info_cell)
        else:
            layout.columns[0].width = Cm(17.0)
            info_cell = layout.rows[0].cells[0]
            photo_cell = info_cell
            cells = (info_cell,)
        for cell in cells:
            self._set_cell_border_v2(cell, border)
            self._set_cell_vertical_alignment(cell)
            cell.text = ""

        photo_paragraph = photo_cell.paragraphs[0]
        photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        photo_paragraph.paragraph_format.space_before = Pt(18)
        self._add_run(photo_paragraph, "◎", size=26, bold=True, color="CBD5E1")
        photo_paragraph.add_run().add_break()
        self._add_run(photo_paragraph, "职业寸照", size=9, bold=True, color="6B7280")
        photo_paragraph.add_run().add_break()
        self._add_run(photo_paragraph, "Professional Photo", size=8.5, color="6B7280")

        table = info_cell.add_table(rows=0, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(3.3)
        table.columns[1].width = Cm(10.1)
        for label, value in rows:
            row = table.add_row()
            for cell in row.cells:
                self._set_cell_border_v2(cell, border)
                self._set_cell_vertical_alignment(cell)
                cell.text = ""
            self._add_run(row.cells[0].paragraphs[0], label, size=8.8, bold=True, color=color)
            self._add_run(row.cells[1].paragraphs[0], value, size=9, color="374151")

    def _doc_personal_info(
        self,
        doc: Document,
        rows: list[tuple[str, str]],
        border: str,
        color: str,
        photo_data_uri: str = "",
        photo_required: bool = False,
    ) -> None:
        show_photo_area = bool(str(photo_data_uri or "").strip()) or bool(photo_required)
        layout = doc.add_table(rows=1, cols=2 if show_photo_area else 1)
        layout.autofit = False
        if show_photo_area:
            layout.columns[0].width = Cm(3.4)
            layout.columns[1].width = Cm(13.6)
            photo_cell, info_cell = layout.rows[0].cells
            cells = (photo_cell, info_cell)
        else:
            layout.columns[0].width = Cm(17.0)
            info_cell = layout.rows[0].cells[0]
            photo_cell = info_cell
            cells = (info_cell,)
        for cell in cells:
            self._set_cell_border_v2(cell, border)
            self._set_cell_vertical_alignment(cell)
            cell.text = ""

        photo_paragraph = photo_cell.paragraphs[0]
        photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if self._add_photo_to_paragraph(photo_paragraph, photo_data_uri):
            photo_paragraph.paragraph_format.space_before = Pt(3)
            photo_paragraph.paragraph_format.space_after = Pt(3)
        else:
            photo_paragraph.paragraph_format.space_before = Pt(18)
            self._add_run(photo_paragraph, "◎", size=26, bold=True, color="CBD5E1")
            photo_paragraph.add_run().add_break()
            self._add_run(photo_paragraph, "职业寸照", size=9, bold=True, color="6B7280")
            photo_paragraph.add_run().add_break()
            self._add_run(photo_paragraph, "Professional Photo", size=8.5, color="6B7280")

        if not show_photo_area:
            info_cell.text = ""

        table = info_cell.add_table(rows=0, cols=4)
        table.autofit = False
        table.columns[0].width = Cm(2.7 if show_photo_area else 3.0)
        table.columns[1].width = Cm(4.0 if show_photo_area else 5.3)
        table.columns[2].width = Cm(2.7 if show_photo_area else 3.0)
        table.columns[3].width = Cm(4.0 if show_photo_area else 5.3)
        pairs = list(rows)
        for index in range(0, len(pairs), 2):
            row = table.add_row()
            for cell in row.cells:
                self._set_cell_border_v2(cell, border)
                self._set_cell_vertical_alignment(cell)
                cell.text = ""
            left = pairs[index]
            right = pairs[index + 1] if index + 1 < len(pairs) else ("", "")
            self._add_run(row.cells[0].paragraphs[0], left[0], size=8.8, bold=True, color=color)
            self._add_run(row.cells[1].paragraphs[0], left[1], size=9, color="374151")
            if right[0]:
                self._add_run(row.cells[2].paragraphs[0], right[0], size=8.8, bold=True, color=color)
                self._add_run(row.cells[3].paragraphs[0], right[1], size=9, color="374151")

    def _add_photo_to_paragraph(self, paragraph: Any, data_uri: str) -> bool:
        match = re.match(r"^data:image/(?:png|jpe?g|webp);base64,([A-Za-z0-9+/=\s]+)$", str(data_uri or ""), re.IGNORECASE)
        if not match:
            return False
        try:
            image_bytes = base64.b64decode(re.sub(r"\s+", "", match.group(1)), validate=True)
            paragraph.add_run().add_picture(BytesIO(image_bytes), width=Cm(2.8))
            return True
        except Exception:
            return False

    def _set_cell_border_v2(self, cell: Any, color: str) -> None:
        """Lighter border for a cleaner look."""
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            element = borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "2")
            element.set(qn("w:color"), color.replace("#", ""))

    def _resume_experience_section_v2(self, doc: Document, data: dict[str, Any], color: str) -> None:
        """Improved experience section with better bullet formatting."""
        resume_text = self._resume_text(data)
        if not resume_text:
            p = doc.add_paragraph()
            self._add_run(p, "No parsed resume text was found. Please re-upload the resume file.", size=9.5, color="9CA3AF")
            return

        sections = resume_sections_from_data(data)
        if sections:
            titles = {
                "summary": "Profile / \u81ea\u6211\u8bc4\u4ef7",
                "experience": "Work Experience / \u5de5\u4f5c\u7ecf\u5386",
                "skills": "Skills / \u6838\u5fc3\u6280\u80fd",
                "education": "Education / \u6559\u80b2\u7ecf\u5386",
                "projects": "Project Experience / \u9879\u76ee\u7ecf\u5386",
                "certificates": "Certificates / \u8363\u8a89\u8bc1\u4e66",
                "unclassified": "Parsed Resume / \u7b80\u5386\u539f\u6587",
            }
            for key in ["summary", "experience", "projects", "skills", "education", "certificates", "unclassified"]:
                items = sections.get(key) or []
                if not items:
                    continue
                heading = doc.add_paragraph()
                heading.paragraph_format.space_before = Pt(6)
                heading.paragraph_format.space_after = Pt(3)
                self._add_run(heading, titles.get(key, key.title()), size=10, bold=True, color=color)
                for item in items[:10]:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.45)
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.line_spacing = Pt(14)
                    self._add_run(p, item, size=9.2, color="374151")
            return

        bullets = self._resume_bullets(resume_text)
        if not bullets:
            bullets = [resume_text[:200].strip()]
            
        for bullet in bullets:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = Pt(16)
            self._add_run(p, "• ", size=9.5, bold=True, color=color)
            self._add_run(p, bullet, size=9.5, color="374151")

    def _doc_bullets(self, doc: Document, items: list[str], color: str) -> None:
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.55)
            p.paragraph_format.first_line_indent = Cm(-0.35)
            p.paragraph_format.space_after = Pt(3)
            self._add_run(p, "\u2022 ", size=9.3, bold=True, color=color)
            self._add_run(p, str(item), size=9.3, color="374151")

    def _doc_experience_groups(self, doc: Document, groups: list[dict[str, Any]], color: str) -> None:
        for group in groups:
            company = doc.add_paragraph()
            company.paragraph_format.space_before = Pt(5)
            company.paragraph_format.space_after = Pt(2)
            self._add_run(company, str(group.get("company") or "工作经历"), size=10, bold=True, color="111827")
            for role in group.get("roles", []):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.35)
                p.paragraph_format.space_after = Pt(2)
                self._add_run(p, str(role.get("period") or "-"), size=9.3, bold=True, color=color)
                title = str(role.get("title") or "").strip()
                if title:
                    self._add_run(p, f"  {title}", size=9.3, bold=True, color="374151")
                for detail in role.get("details", [])[:5]:
                    dp = doc.add_paragraph()
                    dp.paragraph_format.left_indent = Cm(0.75)
                    dp.paragraph_format.first_line_indent = Cm(-0.3)
                    dp.paragraph_format.space_after = Pt(2)
                    self._add_run(dp, "\u2022 ", size=9, bold=True, color=color)
                    self._add_run(dp, str(detail), size=9, color="374151")

    def _resume_section_blocks(self, doc: Document, blocks: list[dict[str, Any]], color: str) -> None:
        for block in blocks:
            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(5)
            heading.paragraph_format.space_after = Pt(2)
            self._add_run(heading, str(block.get("title") or ""), size=10, bold=True, color=color)
            for item in block.get("items", [])[:10]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.35)
                p.paragraph_format.space_after = Pt(2)
                self._add_run(p, str(item), size=9.1, color="374151")

    def _original_resume_appendix(self, doc: Document, text: str) -> None:
        value = str(text or "").strip()
        if not value:
            value = "No original resume text was found. Please re-upload the resume file."
        for raw_line in value.splitlines():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = Pt(13)
            self._add_run(p, raw_line if raw_line.strip() else " ", size=8.8, color="111827")

    def _text_box_v2(self, doc: Document, text: str, border: str) -> None:
        """Improved text box with lighter border."""
        table = doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = Cm(17.0)
        panel = table.rows[0].cells[0]
        panel.text = ""
        self._set_cell_border_v2(panel, border)
        self._add_run(panel.paragraphs[0], text or "-", size=9.5, color="374151")

    def _add_info_table(self, doc: Document, data: dict[str, Any], fields: list[str]) -> None:
        """添加信息表格"""
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        for field_name in fields:
            value = data.get(field_name, "")
            if not value:
                continue

            row = table.add_row()
            row.cells[0].text = field_name.replace("_", " ").title()
            row.cells[1].text = str(value)

        doc.add_paragraph()

    def _add_assessment_section(self, doc: Document, data: dict[str, Any]) -> None:
        """添加顾问评估部分"""
        assessment_fields = [
            ("motivation", "跳槽动因"),
            ("recommendation_rationale", "推荐理由"),
            ("opportunity_to_improve", "可提升空间"),
            ("role_fit", "岗位契合度"),
        ]

        for field_name, label in assessment_fields:
            value = data.get(field_name)
            if not value:
                continue

            doc.add_heading(label, level=2)

            if isinstance(value, dict):
                for sub_key, sub_value in value.items():
                    if sub_value:
                        p = doc.add_paragraph()
                        p.add_run(f"{sub_key.replace('_', ' ').title()}: ").bold = True
                        p.add_run(str(sub_value))
            else:
                doc.add_paragraph(str(value))

    # ============================================================
    # 工具方法
    # ============================================================

    def _set_cell_text(self, cell: Any, text: str) -> None:
        """设置单元格文本并应用字体"""
        cell.text = text
        self._apply_font_to_cell(cell)

    def _apply_font_to_cell(self, cell: Any, bold: bool = False, italic: bool = False, color: str | None = None) -> None:
        """应用字体样式到单元格"""
        font_family = self.font_config.get("family", "Microsoft YaHei")
        font_family_en = self.font_config.get("family_en", "Arial")
        font_size = self.font_config.get("size", 10.5)

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_family_en
                run.font.size = Pt(font_size)
                run.bold = bold
                run.italic = italic
                if color:
                    run.font.color.rgb = RGBColor.from_string(color)

                # 设置中文字体
                run._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)

    def _apply_global_styles(self, doc: Document) -> None:
        """应用全局样式"""
        # 遍历所有段落和表格单元格，统一字体
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                self._apply_font_to_run(run)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            self._apply_font_to_run(run)

    def _apply_document_font(self, doc: Document) -> None:
        """应用文档默认字体"""
        style = doc.styles["Normal"]
        style.font.name = self.font_config.get("family_en", "Arial")
        style.font.size = Pt(self.font_config.get("size", 10.5))
        style._element.rPr.rFonts.set(qn("w:eastAsia"), self.font_config.get("family", "Microsoft YaHei"))

    def _apply_font_to_run(self, run: Any) -> None:
        """应用字体到 run"""
        font_family = self.font_config.get("family", "Microsoft YaHei")
        font_family_en = self.font_config.get("family_en", "Arial")
        font_size = self.font_config.get("size", 10.5)

        run.font.name = font_family_en
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)

    def _add_run(
        self,
        paragraph: Any,
        text: str,
        *,
        size: float | None = None,
        bold: bool = False,
        color: str | None = None,
    ) -> Any:
        """添加格式化的文本 run"""
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = self.font_config.get("family_en", "Arial")
        run.font.size = Pt(size or self.font_config.get("size", 10.5))
        if color:
            run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.font_config.get("family", "Microsoft YaHei"))
        return run

    def _shade_cell(self, cell: Any, fill: str) -> None:
        """设置单元格背景色"""
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), fill.replace("#", ""))

    def _set_cell_border(self, cell: Any, color: str) -> None:
        """设置单元格边框"""
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            element = borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:color"), color.replace("#", ""))

    def _remove_table_borders(self, table: Any) -> None:
        """移除表格所有边框"""
        borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            table._tbl.tblPr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            element.set(qn("w:val"), "nil")

    def _set_cell_vertical_alignment(self, cell: Any) -> None:
        """设置单元格垂直对齐"""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _brand_color(self, key: str, fallback: str) -> str:
        """获取品牌颜色"""
        return str(self.brand_config.get("branding", {}).get(key, fallback)).replace("#", "").upper()

    def _value(self, data: dict[str, Any], key: str, fallback: str = "-") -> str:
        """获取数据值"""
        value = data.get(key)
        return fallback if value is None or value == "" else str(value)

    def _nested_value(self, data: Any, key: str, fallback: str = "-") -> str:
        """获取嵌套字典值"""
        return str(data[key]) if isinstance(data, dict) and data.get(key) else fallback

    def _resume_text(self, data: dict[str, Any]) -> str:
        """获取简历文本"""
        return str(data.get("original_resume") or data.get("resume_text") or "").strip()

    def _resume_bullets(self, text: str) -> list[str]:
        """将简历文本分割为 bullet points"""
        cleaned = re.sub(r"\s+", " ", text).strip()
        if not cleaned:
            return []
        parts = re.split(r"(?<=[.!?\u3002\uff01\uff1f\uff1b;])\s*", cleaned)
        bullets = [item.strip(" -\u2022\t") for item in parts if len(item.strip(" -\u2022\t")) >= 8]
        if len(bullets) < 2:
            bullets = [cleaned[i:i + 160].strip() for i in range(0, min(len(cleaned), 520), 160)]
        return [self._truncate(item, 220) for item in bullets[:4] if item]

    def _compact_join(self, values: list[Any]) -> str:
        """紧凑连接多个值"""
        return " / ".join(str(value) for value in values if value)

    def _truncate(self, text: str, limit: int) -> str:
        """截断文本"""
        normalized = re.sub(r"\s+", " ", text).strip()
        return normalized if len(normalized) <= limit else normalized[:limit].rstrip() + "..."

    def _get_nested_value(self, data: dict[str, Any], path: str) -> Any:
        """获取嵌套字典值"""
        parts = path.split(".")
        current = data

        for part in parts:
            if isinstance(current, dict) and part in current:
                current = current[part]
            else:
                return None

        return current

    def _format_value(self, value: Any, format_config: dict[str, Any] | None) -> str:
        """根据格式配置格式化值"""
        if not format_config:
            return str(value)

        fmt_type = format_config.get("type")

        if fmt_type == "date" and isinstance(value, str):
            # 简单日期格式化
            from datetime import datetime
            try:
                dt = datetime.strptime(value, "%Y-%m-%d")
                pattern = format_config.get("pattern", "%Y-%m-%d")
                return dt.strftime(pattern)
            except ValueError:
                return str(value)

        if fmt_type == "number" and isinstance(value, (int, float)):
            suffix = format_config.get("suffix", "")
            return f"{value}{suffix}"

        return str(value)


# 便捷函数
__all__ = ["ReportRenderer"]
