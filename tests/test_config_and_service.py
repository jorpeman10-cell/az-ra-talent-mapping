import tempfile
import unittest
import json
import zipfile
from pathlib import Path

from core.config_loader import ConfigValidator, get_loader


PROJECT_ROOT = Path(__file__).resolve().parents[1]


class ConfigAndServiceTests(unittest.TestCase):
    def test_default_brand_loads_and_validates(self):
        loader = get_loader(PROJECT_ROOT / "config")
        brand = loader.load_brand("default")
        self.assertEqual(brand["brand_id"], "default")
        self.assertTrue(brand["fields"]["required"])
        validator = ConfigValidator(loader)
        self.assertTrue(validator.validate_brand("default"), validator.errors)

    def test_hiijob_agent_fallback_generates_deterministic_comments(self):
        from core.hiijob_agent import HiijobAgentClient

        client = HiijobAgentClient(base_url="", agent_id="", token="")
        result = client.generate_comments({
            "candidate_name": "Alice",
            "position_title": "Medical Director",
            "resume_text": "Alice has oncology launch and medical affairs experience.",
            "job_description": "Lead oncology medical strategy and cross-functional launch planning.",
        })

        comments = result["comments"]
        self.assertIn("recommendation_rationale", comments)
        self.assertIn("Medical Director", comments["role_fit"])
        self.assertIn("JD", comments["recommendation_rationale"]["strengths_summary"])
        self.assertIn("resume", comments["recommendation_rationale"]["strengths_summary"].lower())
        self.assertIn("oncology", comments["recommendation_rationale"]["strengths_summary"])
        self.assertNotIn("CN:", comments["role_fit"])
        self.assertNotIn("EN:", comments["role_fit"])
        self.assertTrue(result["missing_information"])

    def test_hiijob_agent_fallback_generates_chinese_comments_without_bilingual_mix(self):
        from core.hiijob_agent import HiijobAgentClient

        client = HiijobAgentClient(base_url="", agent_id="", token="")
        result = client.generate_comments({
            "candidate_name": "何超人",
            "position_title": "RPM",
            "resume_text": "阿斯利康 DSM，负责 ZOK 推广，销售额增长。",
            "job_description": "负责区域市场准入、团队管理与核心客户维护。",
        })

        comments = result["comments"]
        combined = " ".join([
            comments["recommendation_rationale"]["strengths_summary"],
            comments["recommendation_rationale"]["risk_notes"],
            comments["motivation"],
            comments["role_fit"],
        ])
        self.assertIn("何超人", combined)
        self.assertIn("RPM", combined)
        self.assertIn("简历", combined)
        self.assertNotIn("EN:", combined)
        self.assertNotIn("CN:", combined)

    def test_report_service_creates_draft_and_card_url(self):
        from core.report_service import ReportService

        with tempfile.TemporaryDirectory() as tmp:
            service = ReportService(
                config_dir=PROJECT_ROOT / "config",
                data_dir=Path(tmp),
                public_base_url="http://testserver",
            )
            draft = service.create_draft({
                "brand_id": "default",
                "candidate_name": "Alice",
                "position_title": "Medical Director",
                "resume_text": "Oncology medical affairs leader.",
            })

            self.assertEqual(draft["status"], "draft")
            self.assertTrue(draft["report_id"].startswith("report_"))
            self.assertIn("/cards/reports/", draft["card_url"])
            self.assertEqual(draft["data"]["resume_quality"]["status"], "low_confidence")

    def test_report_service_flags_low_quality_resume_text(self):
        from core.report_service import ReportService

        with tempfile.TemporaryDirectory() as tmp:
            service = ReportService(
                config_dir=PROJECT_ROOT / "config",
                data_dir=Path(tmp),
                public_base_url="http://testserver",
            )
            draft = service.create_draft({
                "brand_id": "tstar",
                "candidate_name": "Alice",
                "position_title": "CRA",
                "resume_text": "",
            })

            quality = draft["data"].get("resume_quality", {})
            self.assertEqual(quality["status"], "needs_ocr")
            self.assertTrue(quality["needs_ocr"])

    def test_report_service_materializes_resume_source_and_candidate_brief(self):
        from core.report_service import ReportService

        resume = "\n".join([
            "个人信息",
            "姓名：何超人 电话：13732237830 邮箱：2294338095@qq.com",
            "工作经历",
            "2026.01-2026.06 阿斯利康 DSM 负责萧山区核心市场推广。",
            "2025.01-2025.12 EPS 负责产品准入和客户维护。",
            "FULL_RESUME_SOURCE_TAIL_MARKER",
        ])
        with tempfile.TemporaryDirectory() as tmp:
            data_dir = Path(tmp)
            service = ReportService(
                config_dir=PROJECT_ROOT / "config",
                data_dir=data_dir,
                public_base_url="http://testserver",
            )

            draft = service.create_draft({
                "brand_id": "tstar",
                "candidate_name": "何超人",
                "position_title": "RPM",
                "resume_text": resume,
                "salary_info": "Total 50w",
            })

            data = draft["data"]
            self.assertRegex(data["resume_source_id"], r"^rs_[a-f0-9]{32}$")
            self.assertRegex(data["candidate_brief_id"], r"^cb_[a-f0-9]{32}$")
            self.assertIn("candidate_brief", data)
            self.assertEqual(data["candidate_brief"]["resume_source_id"], data["resume_source_id"])
            self.assertIn("career_history", data["candidate_brief"])
            self.assertTrue(data["candidate_brief"]["career_history"])

            source_path = data_dir / "resume_sources" / f"{data['resume_source_id']}.json"
            self.assertTrue(source_path.exists())
            source_payload = json.loads(source_path.read_text(encoding="utf-8"))
            self.assertIn("FULL_RESUME_SOURCE_TAIL_MARKER", source_payload["text"])

            reference_draft = service.create_draft({
                "brand_id": "tstar",
                "candidate_name": "何超人",
                "position_title": "RPM",
                "candidate_brief_id": data["candidate_brief_id"],
            })

            self.assertIn("FULL_RESUME_SOURCE_TAIL_MARKER", reference_draft["data"]["original_resume"])
            self.assertEqual(reference_draft["data"]["resume_source_id"], data["resume_source_id"])
            self.assertEqual(reference_draft["data"]["candidate_brief_id"], data["candidate_brief_id"])

    def test_report_service_default_public_url_uses_b9_port(self):
        from core.report_service import ReportService

        with tempfile.TemporaryDirectory() as tmp:
            service = ReportService(config_dir=PROJECT_ROOT / "config", data_dir=Path(tmp))
            draft = service.create_draft({
                "brand_id": "default",
                "candidate_name": "Alice",
                "position_title": "Medical Director",
            })

            self.assertTrue(draft["card_url"].startswith("http://localhost:8810/"))

    def test_report_service_renders_docx(self):
        from core.report_service import ReportService

        with tempfile.TemporaryDirectory() as tmp:
            service = ReportService(
                config_dir=PROJECT_ROOT / "config",
                data_dir=Path(tmp),
                public_base_url="http://testserver",
            )
            draft = service.create_draft({
                "brand_id": "default",
                "candidate_name": "Alice",
                "position_title": "Medical Director",
            })
            service.generate_comments(draft["report_id"])
            rendered = service.render_report(draft["report_id"])

            self.assertEqual(rendered["status"], "confirmed")
            self.assertTrue((Path(tmp) / "outputs" / rendered["filename"]).exists())

    def test_tstar_docx_uses_bilingual_template_structure(self):
        from docx import Document

        from core.report_service import ReportService

        with tempfile.TemporaryDirectory() as tmp:
            service = ReportService(
                config_dir=PROJECT_ROOT / "config",
                data_dir=Path(tmp),
                public_base_url="http://testserver",
            )
            draft = service.create_draft({
                "brand_id": "tstar",
                "candidate_name": "ZHENG Huang",
                "position_title": "地区经理",
                "salary_info": "20K*12 with annual bonus",
                "resume_text": "个人信息\n姓名：ZHENG Huang\n工作经历\n2020.06 - 至今\n阿斯利康医药（杭州）有限公司\n保留率在65%以上。",
                "job_description": "负责区域销售团队与核心医院准入。",
            })
            service.generate_comments(draft["report_id"])
            rendered = service.render_report(draft["report_id"])
            output_path = Path(tmp) / "outputs" / rendered["filename"]
            doc = Document(str(output_path))
            with zipfile.ZipFile(output_path) as package:
                document_xml = package.read("word/document.xml").decode("utf-8")
            paragraph_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            table_text = "\n".join(
                cell.text
                for table in doc.tables
                for row in table.rows
                for cell in row.cells
            )
            full_text = paragraph_text + "\n" + table_text
            footer_text = doc.sections[0].footer.paragraphs[0].text

            self.assertIn("Candidate Profile", full_text)
            self.assertIn("Work Experience", full_text)
            self.assertIn("Original Resume Appendix", full_text)
            self.assertNotIn("Parsing Confidence", full_text)
            self.assertNotIn("Structured Resume", full_text)
            self.assertNotIn("封面与使用说明", full_text)
            self.assertIn("20K*12 with annual bonus", document_xml)
            self.assertIn('<w:color w:val="111827"/>', document_xml)
            self.assertNotIn("职业寸照", full_text)
            self.assertNotIn("Professional Photo", full_text)
            self.assertIn("2020.06", full_text)
            self.assertIn("Work Experience", full_text)
            self.assertIn("Appendix:", footer_text)
            self.assertGreaterEqual(len(doc.tables), 5)

    def test_report_service_renders_review_card_html(self):
        from core.report_service import ReportService

        with tempfile.TemporaryDirectory() as tmp:
            service = ReportService(
                config_dir=PROJECT_ROOT / "config",
                data_dir=Path(tmp),
                public_base_url="http://testserver",
            )
            draft = service.create_draft({
                "brand_id": "default",
                "candidate_name": "Alice",
                "position_title": "Medical Director",
            })
            html = service.render_card_html(draft["report_id"])

            self.assertIn("Alice", html)
            self.assertIn("Medical Director", html)
            self.assertIn("Generate DOCX", html)
            self.assertIn("DOCX Template Diagnostics", html)

    def test_report_service_refreshes_stale_parsed_resume_before_html_render(self):
        from core.report_service import ReportService

        resume = (
            "个人信息\n"
            "姓名：杨炯铭 电话：13625816396\n"
            "工作经历\n"
            "2020.6月--至今 阿斯利康医药（杭州）有限公司 高级地区经理 "
            "负责区域团队管理，市场保留率在65%以上。\n"
            "2019.02-2020.06 阿斯利康医药（杭州）有限公司 S,MR "
            "负责血脂康产品推广。\n"
            "教育经历\n"
            "浙江工业大学 本科（药学专业）\n"
        )
        with tempfile.TemporaryDirectory() as tmp:
            data_dir = Path(tmp)
            service = ReportService(
                config_dir=PROJECT_ROOT / "config",
                data_dir=data_dir,
                public_base_url="http://testserver",
            )
            draft = service.create_draft({
                "brand_id": "tstar",
                "candidate_name": "杨炯铭",
                "position_title": "地区经理",
                "resume_text": resume,
            })
            record_path = data_dir / "drafts" / f"{draft['report_id']}.json"
            record = json.loads(record_path.read_text(encoding="utf-8"))
            record["data"]["parsed_resume"] = {
                "text": resume,
                "structured": {
                    "sections": {
                        "experience": ["学历：本科（药学专业）", "住址：浙江省杭州市上城区"],
                    },
                    "experience_items": ["学历：本科（药学专业）"],
                },
            }
            record_path.write_text(json.dumps(record, ensure_ascii=False), encoding="utf-8")

            rendered = service.render_html_report(draft["report_id"])
            html_text = (data_dir / "outputs" / rendered["filename"]).read_text(encoding="utf-8")

            self.assertIn("阿斯利康医药", html_text)
            self.assertIn("高级地区经理", html_text)
            self.assertIn("S,MR", html_text)

    def test_html_report_removes_visible_bullet_markers(self):
        from core.report_service import ReportService

        resume = (
            "工作经历\n"
            "2020.06-至今 阿斯利康医药（杭州）有限公司 高级地区经理\n"
            "• 负责区域团队管理。\n"
            "•\n"
            "• 达成93%。\n"
        )
        with tempfile.TemporaryDirectory() as tmp:
            data_dir = Path(tmp)
            service = ReportService(
                config_dir=PROJECT_ROOT / "config",
                data_dir=data_dir,
                public_base_url="http://testserver",
            )
            draft = service.create_draft({
                "brand_id": "tstar",
                "candidate_name": "杨炯铭",
                "position_title": "地区经理",
                "resume_text": resume,
                "recommendation_rationale": {
                    "strengths_summary": "• 负责区域团队管理，达成93%。",
                    "risk_notes": "• 待确认到岗时间。",
                },
            })

            rendered = service.render_html_report(draft["report_id"])
            html_text = (data_dir / "outputs" / rendered["filename"]).read_text(encoding="utf-8")

            self.assertIn("list-style: none", html_text)
            self.assertNotIn(">•", html_text)

    def test_html_report_omits_style_badge_and_does_not_duplicate_salary(self):
        from core.html_renderer import render_report_html

        html_text = render_report_html(
            {
                "brand_id": "tstar",
                "candidate_name": "ZHENG Huang",
                "position_title": "地区经理",
                "salary_info": "20K*12 with annual bonus",
                "resume_text": "个人信息\n姓名：ZHENG Huang\n出生年月：1988.12\n工作经历\n2020.06-至今 测试有限公司 地区经理",
            },
            {"brand_id": "tstar"},
        )

        self.assertNotIn("T-STAR Burgundy", html_text)
        self.assertEqual(html_text.count("20K*12 with annual bonus"), 1)
        self.assertIn("candidate-profile-layout no-photo", html_text)
        self.assertNotIn('<div class="professional-photo-placeholder"', html_text)
        self.assertNotIn('<span class="photo-avatar-icon"', html_text)
        self.assertNotIn("Professional Photo", html_text)

    def test_html_report_shows_photo_placeholder_only_when_requested(self):
        from core.html_renderer import render_report_html

        html_text = render_report_html(
            {
                "brand_id": "tstar",
                "candidate_name": "ZHENG Huang",
                "position_title": "地区经理",
                "professional_photo_required": True,
                "resume_text": "个人信息\n姓名：ZHENG Huang\n工作经历\n2020.06-至今 测试有限公司 地区经理",
            },
            {"brand_id": "tstar"},
        )

        self.assertIn("candidate-profile-layout has-photo", html_text)
        self.assertIn("professional-photo-placeholder", html_text)
        self.assertIn("Photo pending", html_text)

    def test_html_report_profile_includes_contact_rows_and_uploaded_photo(self):
        from core.html_renderer import render_report_html

        photo_data_uri = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAFgwJ/lZ9vAAAAAABJRU5ErkJggg=="
        html_text = render_report_html(
            {
                "brand_id": "tstar",
                "candidate_name": "何超人",
                "position_title": "RPM",
                "current_title": "DSM",
                "current_company": "AZ",
                "professional_photo_data_uri": photo_data_uri,
                "resume_text": (
                    "个人信息\n"
                    "姓名：何超人 电话：13732237830 邮箱：2294338095@qq.com 毕业院校：浙江中医药大学\n"
                    "工作经历\n"
                    "2026.01-2026.06 阿斯利康 DSM 负责市场推广。"
                ),
            },
            {"brand_id": "tstar"},
        )

        self.assertIn("13732237830", html_text)
        self.assertIn("2294338095@qq.com", html_text)
        self.assertIn("professional-photo", html_text)
        self.assertIn(photo_data_uri, html_text)
        self.assertNotIn('<span class="photo-avatar-icon"', html_text)

    def test_tstar_docx_embeds_uploaded_professional_photo(self):
        from docx import Document

        from core.report_service import ReportService

        photo_data_uri = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAFgwJ/lZ9vAAAAAABJRU5ErkJggg=="
        with tempfile.TemporaryDirectory() as tmp:
            service = ReportService(
                config_dir=PROJECT_ROOT / "config",
                data_dir=Path(tmp),
                public_base_url="http://testserver",
            )
            draft = service.create_draft({
                "brand_id": "tstar",
                "candidate_name": "何超人",
                "position_title": "RPM",
                "professional_photo_data_uri": photo_data_uri,
                "resume_text": "个人信息\n姓名：何超人 电话：13732237830 邮箱：2294338095@qq.com\n工作经历\n2026.01-2026.06 阿斯利康 DSM",
                "recommendation_rationale": {"strengths_summary": "匹配", "risk_notes": "待确认"},
                "motivation": "待确认",
                "role_fit": "初步匹配",
            })
            rendered = service.render_report(draft["report_id"])
            output_path = Path(tmp) / "outputs" / rendered["filename"]
            doc = Document(str(output_path))

            self.assertGreaterEqual(len(doc.inline_shapes), 1)
            full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            self.assertNotIn("Professional Photo", full_text)
            with zipfile.ZipFile(output_path) as package:
                names = package.namelist()
            self.assertTrue(any(name.startswith("word/media/") for name in names))

    def test_tstar_docx_omits_photo_placeholder_by_default(self):
        from docx import Document

        from core.report_service import ReportService

        with tempfile.TemporaryDirectory() as tmp:
            service = ReportService(
                config_dir=PROJECT_ROOT / "config",
                data_dir=Path(tmp),
                public_base_url="http://testserver",
            )
            draft = service.create_draft({
                "brand_id": "tstar",
                "candidate_name": "何超人",
                "position_title": "RPM",
                "resume_text": "个人信息\n姓名：何超人 电话：13732237830 邮箱：2294338095@qq.com\n工作经历\n2026.01-2026.06 阿斯利康 DSM",
                "recommendation_rationale": {"strengths_summary": "匹配", "risk_notes": "待确认"},
                "motivation": "待确认",
                "role_fit": "初步匹配",
            })
            rendered = service.render_report(draft["report_id"])
            output_path = Path(tmp) / "outputs" / rendered["filename"]
            doc = Document(str(output_path))

            full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            self.assertNotIn("Professional Photo", full_text)
            self.assertNotIn("Photo pending", full_text)

    def test_resume_parser_extracts_structured_sections(self):
        from core.resume_parser import parse_resume_for_report

        parsed = parse_resume_for_report(
            "2020-2024 Roche Medical Director led oncology launch and KOL strategy.\n"
            "Achieved 120% target and built regional medical affairs team.\n"
            "Managed cross-functional launch planning with market access and sales teams."
        )

        self.assertIn("structured", parsed)
        self.assertTrue(parsed["structured"]["experience_items"])
        self.assertTrue(parsed["structured"]["achievement_items"])
        self.assertIn("oncology", " ".join(parsed["structured"]["keyword_items"]).lower())
        self.assertEqual(parsed["quality"]["status"], "low_confidence")

    def test_resume_parser_quality_marks_empty_text_as_needing_ocr(self):
        from core.resume_parser import parse_resume_for_report

        parsed = parse_resume_for_report("")

        self.assertEqual(parsed["quality"]["status"], "needs_ocr")
        self.assertTrue(parsed["quality"]["needs_ocr"])
        self.assertIn("empty_text", parsed["quality"]["reasons"])

    def test_resume_parser_preserves_chinese_resume_modules(self):
        from core.resume_parser import parse_resume_for_report

        parsed = parse_resume_for_report(
            "\u4e2a\u4eba\u4fe1\u606f\n"
            "\u59d3\u540d\uff1aZHENG Huang\n"
            "\u6c42\u804c\u610f\u5411\n"
            "\u610f\u5411\u5c97\u4f4d\uff1a\u5730\u533a\u7ecf\u7406\n"
            "\u81ea\u6211\u8bc4\u4ef7\n"
            "\u6df1\u8015\u533b\u836f\u884c\u4e1a21\u5e74\uff0c\u5177\u5907\u533a\u57df\u5e02\u573a\u5f00\u62d3\u3001\u56e2\u961f\u7ba1\u7406\u4e0e\u4ea7\u54c1\u51c6\u5165\u7ecf\u9a8c\u3002\n"
            "\u5de5\u4f5c\u7ecf\u5386\n"
            "2020.06 - \u81f3\u4eca\n"
            "\u963f\u65af\u5229\u5eb7\u533b\u836f\uff08\u676d\u5dde\uff09\u6709\u9650\u516c\u53f8\n"
            "\u9ad8\u7ea7\u5730\u533a\u7ecf\u7406 | \u533b\u836f\u5065\u5eb7\n"
            "\u4fdd\u7559\u7387\u572865%\u4ee5\u4e0a\uff0cH1\u8fbe\u621093%\uff0c\u6cfd\u745e84%\u589e\u957f\u3002\n"
            "\u6559\u80b2\u7ecf\u5386\n"
            "\u5609\u83b8\u5802\u533b\u836f\u533b\u5b66\uff08\u836f\u5b66\u4e13\u4e1a\uff09\n"
        )

        sections = parsed["structured"]["sections"]
        self.assertIn("summary", sections)
        self.assertIn("experience", sections)
        self.assertIn("\u6df1\u8015\u533b\u836f\u884c\u4e1a21\u5e74", " ".join(sections["summary"]))
        self.assertIn("\u963f\u65af\u5229\u5eb7\u533b\u836f", " ".join(sections["experience"]))
        self.assertTrue(any("2020.06" in item and "\u963f\u65af\u5229\u5eb7" in item for item in parsed["structured"]["experience_items"]))

    def test_resume_parser_splits_embedded_personal_fields_from_work_history(self):
        from core.resume_parser import parse_resume_for_report, resume_work_experience_from_data
        from core.placeholder_report import build_placeholder_context

        resume = (
            "2020.6月--至今 阿斯利康医药（杭州）有限公司 高级地区经理 26年安达唐集采，市场保留率在65%以上。 "
            "工作经历 姓名：杨炯铭 民族：汉 电话：13625816396 邮箱：yangjm625@163.com "
            "出生年月：1988.12 毕业院校：浙江工业大学 学历：本科（药学专业） 住址：浙江省杭州市上城区"
        )
        parsed = parse_resume_for_report(resume)
        data = {
            "candidate_name": "杨炯铭",
            "position_title": "地区经理",
            "original_resume": resume,
            "parsed_resume": parsed,
        }
        ctx = build_placeholder_context(data, {"brand_id": "tstar"})

        labels = [label for label, _ in ctx["personal_info_rows"]]
        values = " ".join(value for _, value in ctx["personal_info_rows"])
        self.assertIn("Name / 姓名", labels)
        self.assertIn("Phone / 电话", labels)
        self.assertIn("Email / 邮箱", labels)
        personal = dict(ctx["personal_info_rows"])
        self.assertEqual(personal["Name / 姓名"], "杨炯铭")
        self.assertEqual(personal["Phone / 电话"], "13625816396")
        self.assertEqual(personal["Email / 邮箱"], "yangjm625@163.com")
        self.assertEqual(personal["School / 毕业院校"], "浙江工业大学")
        self.assertIn("13625816396", values)
        self.assertTrue(all("2020.6" not in value for _, value in ctx["personal_info_rows"]))
        self.assertTrue(any("2020.6" in item for item in resume_work_experience_from_data(data)))
        experience_text = " ".join(
            " ".join([str(group["company"])] + [str(role.get("title", ""))] + role.get("details", []))
            for group in ctx["appendix_blocks"]["experience_groups"]
            for role in group["roles"]
        )
        self.assertIn("阿斯利康医药", experience_text)
        self.assertNotIn("13625816396", experience_text)
        self.assertNotIn("yangjm625", experience_text)
        self.assertNotIn("浙江工业大学", experience_text)

    def test_resume_parser_splits_spaced_personal_labels(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = (
            "个人信息\n"
            "毕业院校：浙江工业大学 学 历：本科（药学专业）住 址：浙江省杭州市上城区"
            "姓 名：杨炯铭 民 族：汉电 话：13625816396邮 箱：yangjm625@163.com\n"
            "工作经历\n"
            "2020.6月--至今 阿斯利康医药（杭州）有限公司 高级地区经理 26年安达唐集采，市场保留率在65%以上。\n"
        )
        data = {
            "candidate_name": "杨炯铭",
            "position_title": "地区经理",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        personal = dict(ctx["personal_info_rows"])
        appendix_personal = dict(ctx["appendix_blocks"]["personal"])

        self.assertEqual(personal["Name / 姓名"], "杨炯铭")
        self.assertEqual(personal["Phone / 电话"], "13625816396")
        self.assertEqual(personal["Email / 邮箱"], "yangjm625@163.com")
        self.assertEqual(personal["School / 毕业院校"], "浙江工业大学")
        self.assertEqual(personal["Education / 学历"], "本科（药学专业）")
        self.assertEqual(personal["Address / 地址"], "浙江省杭州市上城区")
        self.assertEqual(appendix_personal["姓名"], "杨炯铭")
        self.assertEqual(appendix_personal["电话"], "13625816396")
        self.assertEqual(appendix_personal["邮箱"], "yangjm625@163.com")

    def test_appendix_groups_multiple_roles_under_same_company_and_dedupes_summary(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = (
            "自我评价\n"
            "能够积极乐观的面对各种难题，始终秉承道阻且长，行则将至的理念。"
            "能够积极乐观的面对各种难题，始终秉承道阻且长，行则将至的理念。\n"
            "工作经历\n"
            "2020.06-至今 阿斯利康医药（杭州）有限公司 高级地区经理\n"
            "负责区域团队管理。\n"
            "2019.02-2020.06 阿斯利康医药（杭州）有限公司 S,MR\n"
            "负责血脂康产品推广。\n"
        )
        data = {
            "candidate_name": "杨炯铭",
            "position_title": "地区经理",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        groups = ctx["appendix_blocks"]["experience_groups"]
        self.assertEqual(len(groups), 1)
        self.assertIn("阿斯利康医药", groups[0]["company"])
        self.assertEqual(len(groups[0]["roles"]), 2)
        self.assertEqual(len(ctx["appendix_blocks"]["summary"]), 1)

    def test_placeholder_context_localizes_comments_and_keeps_salary(self):
        from core.placeholder_report import build_placeholder_context

        ctx = build_placeholder_context(
            {
                "candidate_name": "张三",
                "position_title": "地区经理",
                "resume_text": "个人信息\n姓名：张三\n工作经历\n2020.06-至今 测试公司 地区经理",
                "salary_info": "当前80万，期望100万",
                "motivation": "EN: English motivation. CN: 中文动机。",
                "role_fit": "EN: English fit. CN: 中文匹配。",
                "recommendation_rationale": {
                    "strengths_summary": "EN: English strength. CN: 中文亮点。",
                    "risk_notes": "EN: English risk. CN: 中文风险。",
                },
            },
            {"brand_id": "tstar"},
        )

        self.assertEqual(ctx["report_language"], "zh")
        self.assertEqual(ctx["salary_info"], "当前80万，期望100万")
        self.assertEqual(ctx["motivation"], "中文动机。")
        self.assertNotIn("EN:", ctx["strengths_summary"])
        self.assertNotIn("English", ctx["role_fit"])

    def test_resume_parser_keeps_personal_and_experience_separate_for_dense_chinese_resume(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = (
            "个人信息\n"
            "姓名：杨炯铭 民族：汉 电话：13625816396 邮箱：yangjm625@163.com 出生年月：1988.12 "
            "毕业院校：浙江工业大学 学历：本科（药学专业） 住址：浙江省杭州市上城区\n"
            "自我评价\n"
            "能够积极乐观的面对各种难题，始终秉承道阻且长，行则将至的理念。\n"
            "工作经历\n"
            "2020.6月--至今 阿斯利康医药（杭州）有限公司 高级地区经理 26年安达唐集采，市场保留率在65%以上。 "
            "H1达成93%，新产品安达释增长190%，倍泽瑞84%增长。 "
            "25年实现全产品95%达成的同时，安达唐实现21%的增长。\n"
            "2019.02-2020.06 阿斯利康医药（杭州）有限公司 S,MR 转岗血脂康产品专职代表，"
            "区域内主要医院小营社区卫生服务中心在6个月左右的时间，血脂康从300盒增长到3000盒。\n"
            "2015.02-2019.01 阿斯利康医药（杭州）有限公司 S,MR 负责上城区心血管产品推广，"
            "2015年各个产品做到快速准入，实现105%达成，45%的增长。\n"
        )
        data = {
            "candidate_name": "杨炯铭",
            "position_title": "地区经理",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        personal_text = " ".join(f"{label}:{value}" for label, value in ctx["personal_info_rows"])
        groups = ctx["appendix_blocks"]["experience_groups"]
        experience_text = " ".join(
            " ".join([str(group["company"])] + [str(role.get("title", ""))] + role.get("details", []))
            for group in groups
            for role in group["roles"]
        )

        self.assertIn("杨炯铭", personal_text)
        self.assertIn("13625816396", personal_text)
        self.assertIn("浙江工业大学", personal_text)
        self.assertIn("阿斯利康医药", experience_text)
        self.assertIn("高级地区经理", experience_text)
        self.assertIn("S,MR", experience_text)
        self.assertIn("安达唐", experience_text)
        self.assertNotIn("学历：本科", experience_text)
        self.assertNotIn("住址：浙江省杭州市上城区", experience_text)

    def test_resume_parser_preserves_work_experience_when_education_follows(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = (
            "个人信息\n"
            "姓名：杨炯铭 民族：汉 电话：13625816396 邮箱：yangjm625@163.com "
            "出生年月：1988.12 毕业院校：浙江工业大学 学历：本科（药学专业） 住址：浙江省杭州市上城区\n"
            "自我评价\n"
            "能够积极乐观的面对各种难题，始终秉承道阻且长，行则将至的理念。\n"
            "工作经历\n"
            "2020.6月--至今 阿斯利康医药（杭州）有限公司 高级地区经理 "
            "26年安达唐集采，市场保留率在65%以上。H1达成93%，新产品安达释增长190%，倍泽瑞84%增长。\n"
            "2019.02-2020.06 阿斯利康医药（杭州）有限公司 S,MR "
            "转岗血脂康产品专职代表，区域内主要医院小营社区卫生服务中心在6个月左右的时间，血脂康从300盒增长到3000盒。\n"
            "2015.02-2019.01 阿斯利康医药（杭州）有限公司 S,MR "
            "负责上城区心血管产品推广，2015年各个产品做到快速准入，实现105%达成，45%的增长。\n"
            "教育经历\n"
            "浙江工业大学 本科（药学专业）\n"
        )
        data = {
            "candidate_name": "杨炯铭",
            "position_title": "地区经理",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        personal_text = " ".join(f"{label}:{value}" for label, value in ctx["personal_info_rows"])
        groups = ctx["appendix_blocks"]["experience_groups"]
        experience_text = " ".join(
            " ".join([str(group["company"])] + [str(role.get("title", ""))] + role.get("details", []))
            for group in groups
            for role in group["roles"]
        )

        self.assertIn("13625816396", personal_text)
        self.assertIn("阿斯利康医药", experience_text)
        self.assertIn("高级地区经理", experience_text)
        self.assertIn("S,MR", experience_text)
        self.assertIn("安达唐", experience_text)
        self.assertGreaterEqual(sum(len(group["roles"]) for group in groups), 3, groups)
        self.assertNotIn("浙江工业大学 本科", experience_text)

    def test_resume_parser_recovers_work_history_when_docx_extracts_it_under_profile(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = (
            "个人信息\n"
            "姓名：杨炯铭 民族：汉 电话：13625816396 邮箱：yangjm625@163.com 出生年月：1988.12 "
            "2020.6月--至今 阿斯利康医药（杭州）有限公司 高级地区经理 26年安达唐集采，市场保留率在65%以上。 "
            "⚫ 2020 年 6 月至今负责杭州市萧山区、富阳区心血管产品推广，2022 年 H1 团队获得 TOP SALES TEAMS。 "
            "2019.02-2020.06 阿斯利康医药（杭州）有限公司 S,MR 转岗血脂康产品专职代表。 "
            "2015.02-2019.01 阿斯利康医药（杭州）有限公司 S,MR 负责上城区心血管产品推广。\n"
            "工作经历\n"
            "- 学 历：本科（药学专业）\n"
            "• 住 址：浙江省杭州市上城区\n"
            "自我评价自我评价能够积极乐观的面对各种难题，始终秉承道阻且长，行则将至的理念。"
        )
        data = {
            "candidate_name": "杨炯铭",
            "position_title": "地区经理",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        groups = ctx["appendix_blocks"]["experience_groups"]
        experience_text = " ".join(
            " ".join([str(group["company"])] + [str(role.get("title", ""))] + role.get("details", []))
            for group in groups
            for role in group["roles"]
        )

        self.assertIn("阿斯利康医药", experience_text)
        self.assertIn("高级地区经理", experience_text)
        self.assertIn("S,MR", experience_text)
        self.assertIn("TOP SALES TEAMS", experience_text)
        self.assertNotIn("学历：本科", experience_text)
        self.assertNotIn("住址：浙江省杭州市上城区", experience_text)

    def test_appendix_dedupes_same_company_and_period_roles(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = (
            "工作经历\n"
            "2020.6月--至今 阿斯利康医药（杭州）有限公司 高级地区经理 "
            "26年安达唐集采，市场保留率在65%以上。H1达成93%。\n"
            "2019.02-2020.06 阿斯利康医药（杭州）有限公司 S,MR 转岗血脂康产品专职代表。\n"
            "2020.6月--至今 高级地区经理 "
            "23年安达唐实现65%的增长，倍他乐克集采后第一年保留达到90%以上。\n"
        )
        data = {
            "candidate_name": "杨炯铭",
            "position_title": "地区经理",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        groups = ctx["appendix_blocks"]["experience_groups"]
        az_group = next(group for group in groups if "阿斯利康医药" in group["company"])
        periods = [role["period"] for role in az_group["roles"]]
        senior_roles = [
            role for role in az_group["roles"]
            if role["period"] == "2020.6月--至今" and "高级地区经理" in role.get("title", "")
        ]
        details = " ".join(" ".join(role.get("details", [])) for role in senior_roles)

        self.assertEqual(periods.count("2020.6月--至今"), 1)
        self.assertEqual(len(senior_roles), 1)
        self.assertIn("H1达成93%", details)
        self.assertIn("第一年保留达到90%以上", details)

    def test_appendix_repairs_pdf_split_current_role_period_and_title(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = (
            "工作经历\n"
            "6月--至今 阿斯利康医药（杭州）有限公司\n"
            "- 获得\n"
            "2020. 高级地区经理\n"
            "2020 年 6 月至今负责杭州市萧山区、富阳区心血管产品推广，2022 年 H1 团队 TOP TEAMS。\n"
            "阿斯利康医药（杭州）有限公司\n"
            "2019.02-2020.06 S,MR\n"
            "转岗血脂康产品专职代表。\n"
        )
        data = {
            "candidate_name": "杨炯铭",
            "position_title": "地区经理",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        az_group = next(group for group in ctx["appendix_blocks"]["experience_groups"] if "阿斯利康医药" in group["company"])
        current_role = next(role for role in az_group["roles"] if role["period"] == "2020.6月--至今")

        self.assertEqual(current_role["title"], "高级地区经理")
        self.assertIn("萧山区", " ".join(current_role["details"]))
        self.assertNotEqual(current_role["title"], "获得")
        self.assertTrue(any(role["period"] == "2019.02-2020.06" for role in az_group["roles"]))

    def test_appendix_keeps_english_resume_fragments_as_details_not_titles(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = (
            "Work Experience\n"
            "- and consumables.\n"
            "20+ yrs working with MNCs in medical and life science devices, instrument\n"
            "10+ yrs working as General Manager or national function.\n"
            "Solid experiences of organization development and distribution management.\n"
            "Proven track record in strategic thinking, market segmentation, identifying opportunities, growing business and KAM through innovative market-driven\n"
            "2019.04-present TCM National Manager, China\n"
            "Radiometer China\n"
            "Report to: Managing Director, Greater China\n"
            "Main Responsibilities:\n"
            "Lead nationwide sales and application team members to develop our market and achieve business/KPI targets.\n"
            "end-user product training to reactivate IB and inc\n"
            "2017.02 - 2019.03 Managing Director, China\n"
            "Sarstedt (Shanghai) Trading Co., Ltd.\n"
            "Report to: President Sales/R&D\n"
            "Provide overall management of the organization, including sales, marketing, personnel, technological resources and assets.\n"
            "and Company\n"
            "2014.07 - 2017.01 Report to: General Manager APAC\n"
            "Mauna Kea Technologies\n"
            "National Sales & Marketing Manager, China\n"
            "Be in charge of sales, marketing, channels and KAM in China.\n"
            "Turnover 2016 inc\n"
            "1990.07 - 1995.07\n"
            "Shanghai Second Medical University\n"
        )
        data = {
            "candidate_name": "ZHENG Huang",
            "position_title": "MD",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        groups = ctx["appendix_blocks"]["experience_groups"]
        titles = [
            str(role.get("title") or "")
            for group in groups
            for role in group["roles"]
        ]
        periods = [
            str(role.get("period") or "")
            for group in groups
            for role in group["roles"]
        ]
        companies = [str(group.get("company") or "") for group in groups]

        self.assertIn("2019.04-present", periods)
        self.assertIn("2017.02 - 2019.03", periods)
        self.assertIn("2014.07 - 2017.01", periods)
        self.assertIn("TCM National Manager, China", titles)
        self.assertIn("Managing Director, China", titles)
        self.assertIn("Radiometer China", companies)
        self.assertIn("Sarstedt (Shanghai) Trading Co., Ltd.", companies)
        self.assertIn("Mauna Kea Technologies", companies)
        self.assertNotIn("and consumables.", titles)
        self.assertNotIn("end-user product training to reactivate IB and inc", titles)
        self.assertNotIn("and Company", titles)
        self.assertNotIn("Turnover 2016 inc", titles)
        self.assertNotIn("New application inc", companies)
        self.assertNotIn("and Company", companies)
        self.assertNotIn("Turnover 2016 inc", companies)

    def test_parser_keeps_long_english_professional_experience(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        lines = ["PROFESSIONAL EXPERIENCES"]
        lines.extend(f"Filler achievement line {idx} with sales and market growth." for idx in range(90))
        lines.extend([
            "2010.04 - 2014.05 Gebrü der Martin GmbH & Co. KG Shanghai Rep. Office",
            "Chief Representative, China",
            "Report to: VP Sales, Germany",
            "2008.03 - 2010.03 BK Medical ApS",
            "Country Manager, China",
            "2000.10 - 2008.03 bioMérieux (Shanghai) Co., Ltd",
            "Project and Sales Manager",
            "1998.10 - 2000.10 Olympus Optical Co., Ltd. Shanghai Office",
            "Manager, Surgical Division, East China",
        ])
        resume = "\n".join(lines)
        data = {
            "candidate_name": "Tong ZHANG",
            "position_title": "MD",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        company_names = [str(group.get("company") or "") for group in ctx["appendix_blocks"]["experience_groups"]]
        titles = [
            str(role.get("title") or "")
            for group in ctx["appendix_blocks"]["experience_groups"]
            for role in group["roles"]
        ]

        self.assertIn("Gebrü der Martin GmbH & Co. KG Shanghai Rep. Office", company_names)
        self.assertIn("BK Medical ApS", company_names)
        self.assertIn("bioMérieux (Shanghai) Co., Ltd", company_names)
        self.assertIn("Olympus Optical Co., Ltd. Shanghai Office", company_names)
        self.assertNotIn("PROFESSIONAL EXPERIENCES", company_names)
        self.assertIn("Chief Representative, China", titles)
        self.assertIn("Country Manager, China", titles)
        self.assertIn("Project and Sales Manager", titles)

    def test_english_self_evaluation_stays_out_of_work_experience(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "SELF-EVALUATION",
            "Background of surgeon, specialized in Clinical Medicine.",
            "20+ yrs working with MNCs in medical and life science devices, instrument and consumables.",
            "10+ yrs working as General Manager or national function.",
            "Solid experiences of organization development and distribution management.",
            "PROFESSIONAL EXPERIENCES",
            "2019.04 –present  Radiometer China",
            "TCM National Manager, China",
            "Report to: Managing Director, Greater China",
        ])
        parsed = parse_resume_for_report(resume)
        data = {
            "candidate_name": "Tom Zhang",
            "position_title": "MD",
            "original_resume": resume,
            "parsed_resume": parsed,
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        summary_text = " ".join(parsed["structured"]["sections"].get("summary", []))
        experience_text = " ".join(parsed["structured"]["sections"].get("experience", []))
        company_names = [str(group.get("company") or "") for group in ctx["appendix_blocks"]["experience_groups"]]

        self.assertIn("20+ yrs working with MNCs", summary_text)
        self.assertNotIn("20+ yrs working with MNCs", experience_text)
        self.assertIn("Radiometer China", company_names)
        self.assertNotIn("PROFESSIONAL EXPERIENCES", company_names)

    def test_untitled_short_role_is_demoted_under_previous_role(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "PROFESSIONAL EXPERIENCES",
            "1998.10 – 2000.10 Olympus Optical Co., Ltd. Shanghai Office",
            "Manager, Surgical Division, East China",
            "Report to: Office Representative",
            "1997.07 – 1998.06",
            "Shanghai Second People’s Hospital",
            "1990.07 – 1995.07 Shanghai Second Medical University",
        ])
        data = {
            "candidate_name": "Tom Zhang",
            "position_title": "MD",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        olympus = next(
            group
            for group in ctx["appendix_blocks"]["experience_groups"]
            if group["company"] == "Olympus Optical Co., Ltd. Shanghai Office"
        )
        role_periods = [str(role.get("period") or "") for role in olympus["roles"]]
        hospital = next(
            group
            for group in ctx["appendix_blocks"]["experience_groups"]
            if "Shanghai Second People" in group["company"]
        )
        hospital_role_periods = [str(role.get("period") or "") for role in hospital["roles"]]

        self.assertIn("1998.10 – 2000.10", role_periods)
        self.assertNotIn("1997.07 – 1998.06", role_periods)
        self.assertIn("1997.07 – 1998.06", hospital_role_periods)

    def test_appendix_parses_suffixless_company_role_line(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = (
            "工作经历\n"
            "负责公司\n"
            "2011.07-2013.9 杭州民生医药 产品研发\n"
            "负责仿制药的研发和生产跟进，主要负责项目培美曲塞冻干粉针\n"
        )
        data = {
            "candidate_name": "杨炯铭",
            "position_title": "地区经理",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        groups = ctx["appendix_blocks"]["experience_groups"]
        company_names = [group["company"] for group in groups]
        minsheng = next(group for group in groups if group["company"] == "杭州民生医药")
        role = minsheng["roles"][0]

        self.assertIn("杭州民生医药", company_names)
        self.assertNotIn("负责公司", company_names)
        self.assertEqual(role["period"], "2011.07-2013.9")
        self.assertEqual(role["title"], "产品研发")
        self.assertIn("培美曲塞冻干粉针", " ".join(role["details"]))

    def test_appendix_merges_suffixless_company_duplicate_and_removes_profile_contamination(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = (
            "工作经历\n"
            "2013.09-2015.01 杭州嘉事堂医药医药科技有限公司 市场推广 经理\n"
            "市场推广经理负责公司在浙一和浙二的产品专业推广。\n"
            "2011.07-2013.9 杭州民生医药\n"
            "产品研发负责仿制药的研发和生产跟进，主要负责项目培美曲塞冻干粉针"
            "工作经历工作经历姓名：杨炯铭民族：汉电话：13625816396邮箱：yangjm625@163.com\n"
            "杭州民生医药\n"
            "2011.07-2013.9 产品研发\n"
        )
        data = {
            "candidate_name": "杨炯铭",
            "position_title": "地区经理",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        groups = ctx["appendix_blocks"]["experience_groups"]
        minsheng_groups = [group for group in groups if group["company"] == "杭州民生医药"]
        other_group_text = " ".join(
            " ".join([str(group["company"])] + [str(role.get("title", ""))] + role.get("details", []))
            for group in groups
            if group["company"] != "杭州民生医药"
            for role in group["roles"]
        )

        self.assertEqual(len(minsheng_groups), 1)
        roles = minsheng_groups[0]["roles"]
        self.assertEqual([role["period"] for role in roles].count("2011.07-2013.9"), 1)
        role = next(role for role in roles if role["period"] == "2011.07-2013.9")
        details = " ".join(role.get("details", []))
        self.assertEqual(role["title"], "产品研发")
        self.assertIn("培美曲塞冻干粉针", details)
        self.assertNotIn("杭州民生医药", other_group_text)
        self.assertNotIn("杨炯铭", details)
        self.assertNotIn("13625816396", details)
        self.assertNotIn("yangjm625", details)

    def test_resume_slogan_does_not_pollute_education_appendix(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = (
            "个人信息\n"
            "姓名：杨炯铭 电话：13625816396 邮箱：yangjm625@163.com 出生年月：1988.12\n"
            "工作经历\n"
            "2020.6月--至今 阿斯利康医药（杭州）有限公司 高级地区经理 "
            "负责杭州市萧山区、富阳区心血管产品推广，H1达成93%。\n"
            "教育经历\n"
            "浙江工业大学 本科（药学专业）\n"
            "细心从每一个小细节开始。\n"
        )
        data = {
            "candidate_name": "杨炯铭",
            "position_title": "地区经理",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        education_text = " ".join(ctx["appendix_blocks"]["education"])
        experience_text = " ".join(
            " ".join([str(group["company"])] + [str(role.get("title", ""))] + role.get("details", []))
            for group in ctx["appendix_blocks"]["experience_groups"]
            for role in group["roles"]
        )

        self.assertIn("浙江工业大学", education_text)
        self.assertNotIn("细心从每一个小细节开始", education_text)
        self.assertIn("阿斯利康医药", experience_text)

    def test_resume_parser_separates_education_profile_and_sorts_recent_work_first(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = (
            "工作经历\n"
            "浙江中医药大学\n"
            "-\n"
            "浙江省富阳中学\n"
            "（在校期间获浙江省优秀毕业生、优秀学生干部、三好学生等）\n"
            "2017.03-2018.06\n"
            "阿斯利康实习生\n"
            "负责浙一 ZOK 推广，病房市场占有率提升。\n"
            "2026.01-2026.06 阿斯利康心血管&糖尿病 DSM\n"
            "负责萧山区城核心市场的推广。\n"
            "2025.01-2025.12 EPS\n"
            "负责产品准入和客户维护。\n"
            "教育经历\n"
            "：何超人\n"
            "：13732237830\n"
            "：2294338095@qq.com\n"
            "：本科\n"
            "专业：生物工程\n"
            "籍贯：浙江杭州\n"
        )
        data = {
            "candidate_name": "何超人",
            "position_title": "RPM",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        personal_text = " ".join(f"{label}:{value}" for label, value in ctx["personal_info_rows"])
        education_text = " ".join(ctx["appendix_blocks"]["education"])
        groups = ctx["appendix_blocks"]["experience_groups"]
        experience_text = " ".join(
            " ".join([str(group["company"])] + [str(role.get("period", "")), str(role.get("title", ""))] + role.get("details", []))
            for group in groups
            for role in group["roles"]
        )
        periods = [
            str(role.get("period") or "")
            for group in groups
            for role in group["roles"]
        ]

        self.assertIn("13732237830", personal_text)
        self.assertIn("2294338095@qq.com", personal_text)
        self.assertIn("浙江中医药大学", education_text)
        self.assertIn("浙江省富阳中学", education_text)
        self.assertIn("阿斯利康", experience_text)
        self.assertNotIn("13732237830", experience_text)
        self.assertNotIn("浙江中医药大学", experience_text)
        self.assertTrue(periods[0].startswith("2026.01"), periods)

    def test_mixed_language_resume_keeps_company_date_and_language_sections(self):
        from core.html_renderer import render_report_html
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "Professional Experience",
            "Novo Nordisk (China)",
            "Beijing",
            "National Retail Sales Director, Diabetes Retail Division (DRD)",
            "Mar",
            "2024 – Present",
            "Strategic Planning & Execution: Formulated mid-to-long term pharmacy channel strategies.",
            "Jiangsu Nhwa Pharmaceutical Co., Ltd.",
            "General Manager, CNS Innovation Business Unit",
            "Jul 2023 – Feb 2024",
            "Business Model Innovation: Established new commercial models.",
            "Pfizer Investment Co. Ltd.",
            "Senior National Key Account Manager - Consumer & Dermatology (C&D) Retail",
            "Jan",
            "2013 –",
            "Mar 2017",
            "Account Strategy: Formulated key account annual plans.",
            "Colgate-Palmolive (Guangzhou) Co., Ltd.",
            "Trainee - Beijing Branch",
            "2000 – Apr 2002",
            "张慧敏",
            "现居地：北京",
            "工作经历",
            "诺和诺德 (Novo Nordisk)",
            "北京",
            "全国零售销售总监，商务及零售事业部",
            "2024.03 – 至今",
            "战略规划：制定药店渠道中长期发展策略。",
            "江苏恩华药业",
            "CNS 创新业务事业部 总经理",
            "2023.07 – 2024.02",
            "新业务孵化：构建全新商业模式。",
            "辉瑞 (Pfizer)",
            "2013.01 – 2017.03 零售 高级全国重点客户经理",
            "制定重点客户发展策略及年度生意计划。",
        ])
        data = {
            "brand_id": "tstar",
            "candidate_name": "Moira Zhang",
            "position_title": "Head of Retail",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        groups = ctx["appendix_blocks"]["experience_groups"]
        group_by_company = {group["company"]: group for group in groups}
        novo_role = group_by_company["Novo Nordisk (China)"]["roles"][0]
        pfizer_role = group_by_company["Pfizer Investment Co. Ltd."]["roles"][0]
        cn_pfizer_role = group_by_company["辉瑞 (Pfizer)"]["roles"][0]
        all_experience_text = " ".join(
            " ".join([group["company"], role.get("period", ""), role.get("title", ""), *role.get("details", [])])
            for group in groups
            for role in group["roles"]
        )
        html_text = render_report_html(data, {"brand_id": "tstar"})

        self.assertEqual(novo_role["period"], "Mar 2024 – Present")
        self.assertEqual(novo_role["title"], "National Retail Sales Director, Diabetes Retail Division (DRD)")
        self.assertEqual(pfizer_role["period"], "Jan 2013 – Mar 2017")
        self.assertEqual(pfizer_role["title"], "Senior National Key Account Manager - Consumer & Dermatology (C&D) Retail")
        self.assertEqual(cn_pfizer_role["title"], "零售 高级全国重点客户经理")
        self.assertIn("English Version", html_text)
        self.assertIn("中文版本", html_text)
        self.assertNotIn("张慧敏", all_experience_text)
        self.assertNotIn("现居地：北京", all_experience_text)

    def test_project_blocks_do_not_pollute_hospital_work_experience(self):
        from core.html_renderer import render_report_html
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "个人信息",
            "邮箱：ziva0909@163.com",
            "工作经历",
            "银诺医药 医学顾问",
            "2025.09-2026.03 制定合规且具学术影响力的营销策略方案。",
            "丁香园 高级医学策划",
            "2020.03-2025.09",
            "【医学策略】输出医学推广策略解决方案。",
            "济宁第一人民医院 住院医师",
            "2018.09-2019.11",
            "【临床诊疗】担任内分泌科住院医师。",
            "【住院事项】分管住院病患床位。",
            "减重书籍科普策划 项目策划",
            "2024.10-2025.06",
            "1.基于客户需求及产品特点，策划并撰写250页患者教育内容。",
            "【项目成果】内容专业度获得客户认可。",
        ])
        data = {
            "brand_id": "tstar",
            "candidate_name": "陈静",
            "position_title": "院外运营经理",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        groups = ctx["appendix_blocks"]["experience_groups"]
        projects = ctx["appendix_blocks"]["projects"]
        hospital = next(group for group in groups if group["company"] == "济宁第一人民医院")
        hospital_text = " ".join(
            " ".join([role.get("period", ""), role.get("title", ""), *role.get("details", [])])
            for role in hospital["roles"]
        )
        html_text = render_report_html(data, {"brand_id": "tstar"})

        self.assertIn("2018.09-2019.11", hospital_text)
        self.assertNotIn("2024.10-2025.06", hospital_text)
        self.assertNotIn("减重书籍科普策划", hospital_text)
        self.assertTrue(any("减重书籍科普策划" in item for item in projects))
        self.assertTrue(any("2024.10-2025.06" in item for item in projects))
        self.assertIn("Project Experience / 项目经历", html_text)
        self.assertIn("减重书籍科普策划", html_text)
        self.assertIn("grid-template-columns: 1fr", html_text)
        self.assertIn("overflow-wrap: anywhere", html_text)

    def test_parser_keeps_chinese_resume_work_history_to_real_companies(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "姓名",
            "意向岗位：地区经理",
            "自我评价",
            "深耕医药行业21年，具备成熟的区域市场开拓、团队管理与产品准入经验，擅长核心医院资源维护与新产品上市推广。",
            "工作经历",
            "2020.6-至今 阿斯利康医药(杭州)有限公司",
            "高级地区经理 | 医药健康",
            "负责杭州区域代谢类产品线的整体业务管理，制定区域销售策略，带领12人销售团队完成业绩目标，对接核心三甲医院、基层医疗渠道。",
            "集采应对与市场留存",
            "主导核心产品安达唐集采落地后的市场维护，实现市场保留率稳定在65%以上。",
            "建立重点医院定期沟通机制，协同医学事务部开展临床路径推广。",
            "业绩达成与新产品增长",
            "2026年上半年带领团队完成全年业绩目标的93%。",
            "负责新产品安达释、倍泽瑞的区域上市推广，安达释上市1年销售额同比增长190%。",
            "2012.3-2020.5 赛诺菲(杭州)制药有限公司",
            "地区经理 | 医药健康",
            "负责杭州市上城区心血管产品线的团队管理与业务拓展，对接12家核心三甲医院、社区卫生服务中心。",
            "2015年推动核心心血管产品快速完成区域医院准入，全年实现105%的业绩达成率。",
            "核心技能",
            "销售团队管理：具备10年以上销售团队带教经验。",
            "教育经历",
            "2002.9-2006.6 浙江医科大学",
            "电话：13867129945",
            "邮箱：zhenghuang@outlook.com",
        ])
        data = {
            "candidate_name": "何超人",
            "position_title": "RPM",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        groups = ctx["appendix_blocks"]["experience_groups"]
        company_names = [str(group.get("company") or "") for group in groups]
        experience_text = " ".join(
            " ".join([str(group["company"]), str(role.get("period", "")), str(role.get("title", "")), *role.get("details", [])])
            for group in groups
            for role in group["roles"]
        )

        self.assertEqual(company_names, ["阿斯利康医药(杭州)有限公司", "赛诺菲(杭州)制药有限公司"])
        self.assertIn("高级地区经理", experience_text)
        self.assertIn("地区经理", experience_text)
        self.assertIn("核心三甲医院", experience_text)
        self.assertIn("重点医院定期沟通机制", experience_text)
        self.assertIn("区域医院准入", experience_text)
        self.assertNotIn("核心医院", company_names)
        self.assertNotIn("对接核心三甲医院", company_names)
        self.assertNotIn("2015年推动核心心血管产品快速完成区域医院", company_names)
        self.assertNotIn("销售团队管理", experience_text)
        self.assertNotIn("浙江医科大学", experience_text)

    def test_research_resume_sections_are_preserved_as_projects_not_fake_companies(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "教育背景",
            "学习及科研经历",
            "自我评价",
            "参与项目",
            "专业技能",
            "PD-1 肿瘤治疗等免疫学专业知识",
            "科研经历",
            "研究领域主要为 CREB 介导的信号通路对学习记忆的影响，及其介导药物依赖性的机制以及作为中枢神经药物潜在治疗靶点的研究。",
            "该课题内容发表一作 SCI 论文一篇，影响因子 2.9；国际会议论文摘要一篇，影响因子 1.852。",
            "参与肿瘤治疗免疫检查点研究课题，主要为 PD-1 抑制剂在非小细胞肺癌以及 B 淋巴细胞治疗效果的回顾总结，著有综述。",
            "2013.09 – 2017.06",
            "编写教案和校内练习用书；两班每次月考成绩均在14个班中前两名。",
        ])
        data = {
            "candidate_name": "Research Candidate",
            "position_title": "MSL",
            "original_resume": resume,
            "parsed_resume": parse_resume_for_report(resume),
        }

        ctx = build_placeholder_context(data, {"brand_id": "tstar"})
        projects = ctx["appendix_blocks"]["projects"]
        company_names = [str(group.get("company") or "") for group in ctx["appendix_blocks"]["experience_groups"]]
        project_text = " ".join(projects)

        self.assertIn("CREB", project_text)
        self.assertIn("PD-1", project_text)
        self.assertIn("SCI", project_text)
        self.assertNotIn("科研经历", company_names)
        self.assertNotIn("参与项目", company_names)

    def test_english_resume_headings_do_not_become_companies(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "CURRICULUM VITAE",
            "SUMMARY OF EXPERIENCE",
            "Dedicated Project assistant with almost 2 years of experience in Jiangsu HengRui Medicine Co.,Ltd.",
            "THERAPEUTIC AREA EXPERTISE:",
            "Oncology phase IV study experience.",
            "EMPLOYMENT HISTORY",
            "2020 to present Director, Quality GCP, VISEN Pharmaceuticals (Shanghai) Co. Ltd.",
            "Set up cQMS framework and annual audit program.",
            "16 May 2011 to Feb 2013 Monitor, Clinical operation of PDY, Roche (China) Holding Ltd.",
            "GCP and SOP strictly.",
            "CLINICAL TRIAL EXPERIENCE",
            "Gastric cancer IV 11 sites local study.",
        ])
        ctx = build_placeholder_context(
            {
                "candidate_name": "Test",
                "position_title": "CRA",
                "original_resume": resume,
                "parsed_resume": parse_resume_for_report(resume),
            },
            {"brand_id": "tstar"},
        )
        company_names = [str(group.get("company") or "") for group in ctx["appendix_blocks"]["experience_groups"]]
        project_text = " ".join(ctx["appendix_blocks"]["projects"])

        self.assertIn("VISEN Pharmaceuticals (Shanghai) Co. Ltd.", company_names)
        self.assertIn("Roche (China) Holding Ltd.", company_names)
        self.assertNotIn("EMPLOYMENT HISTORY", company_names)
        self.assertNotIn("CURRICULUM VITAE", company_names)
        self.assertNotIn("THERAPEUTIC AREA EXPERTISE:", company_names)
        self.assertIn("Gastric cancer", project_text)

    def test_hunter_summary_service_company_and_role_become_work_experience(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "简   历",
            "基  本  信  息",
            "移动电话： 13800000000;",
            "公司名称 / 职位",
            "最近工作状态",
            "服务公司： 诺华",
            "行业：",
            "部门职务： PM",
        ])
        ctx = build_placeholder_context(
            {
                "candidate_name": "Test",
                "position_title": "PM",
                "original_resume": resume,
                "parsed_resume": parse_resume_for_report(resume),
            },
            {"brand_id": "tstar"},
        )
        groups = ctx["appendix_blocks"]["experience_groups"]

        self.assertEqual(groups[0]["company"], "诺华")
        self.assertEqual(groups[0]["roles"][0]["title"], "PM")

    def test_english_labelled_employment_blocks_extract_company_title_and_period(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "CURRICULUM VITAE",
            "Employment History",
            "Period: 2019/06 - present",
            "Employer: Everstar Medicines (Shanghai) Limited",
            "Job Title & Function: Senior Project Manager",
            "Main Responsibilities:",
            "Manage and coordinate cross-functional project teams.",
            "Period: 2017/02 - 2019/04",
            "Employer: Novo Nordisk (China) Pharmaceuticals Co., Ltd.",
            "Position: Associate Product Manager, Insulin Brand team",
            "Main Responsibilities:",
            "Support brand strategy development.",
            "Education",
            "China Pharmaceutical University Bachelor",
        ])
        ctx = build_placeholder_context(
            {
                "candidate_name": "Test",
                "position_title": "SPM",
                "original_resume": resume,
                "parsed_resume": parse_resume_for_report(resume),
            },
            {"brand_id": "tstar"},
        )
        groups = ctx["appendix_blocks"]["experience_groups"]
        group_by_company = {group["company"]: group for group in groups}

        self.assertIn("Everstar Medicines (Shanghai) Limited", group_by_company)
        self.assertIn("Novo Nordisk (China) Pharmaceuticals Co., Ltd.", group_by_company)
        everstar_role = group_by_company["Everstar Medicines (Shanghai) Limited"]["roles"][0]
        novo_role = group_by_company["Novo Nordisk (China) Pharmaceuticals Co., Ltd."]["roles"][0]
        self.assertEqual(everstar_role["period"], "2019/06 - present")
        self.assertEqual(everstar_role["title"], "Senior Project Manager")
        self.assertEqual(novo_role["period"], "2017/02 - 2019/04")
        self.assertEqual(novo_role["title"], "Associate Product Manager, Insulin Brand team")
        self.assertNotIn("Main Responsibilities", group_by_company)
        self.assertNotIn("Job Title & Function", group_by_company)

    def test_english_responsibility_sentences_do_not_become_companies(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "Work Experience",
            "RareStone / Shanghai, China",
            "Clinical Project Director",
            "Feb 2022 - Present",
            "Be the clinical operation representative to contribute relevant expertise.",
            "Represent company in cross-functional project meetings.",
            "Timely facilitate the lessons learn and best practice sharing among the PM group.",
            "Syneos Health / Shanghai, China",
            "Project Manager",
            "Oct 2019 - Jan 2022",
            "Liaison with business partners and manage study startup.",
        ])
        ctx = build_placeholder_context(
            {
                "candidate_name": "Test",
                "position_title": "Clinical Project Director",
                "original_resume": resume,
                "parsed_resume": parse_resume_for_report(resume),
            },
            {"brand_id": "tstar"},
        )
        company_names = [str(group.get("company") or "") for group in ctx["appendix_blocks"]["experience_groups"]]
        experience_text = " ".join(
            " ".join([str(group["company"]), role.get("period", ""), role.get("title", ""), *role.get("details", [])])
            for group in ctx["appendix_blocks"]["experience_groups"]
            for role in group["roles"]
        )

        self.assertEqual(company_names, ["RareStone", "Syneos Health"])
        self.assertIn("Clinical Project Director", experience_text)
        self.assertIn("Project Manager", experience_text)
        self.assertIn("Represent company in cross-functional project meetings.", experience_text)
        self.assertNotIn("Represent company", company_names)
        self.assertNotIn("Timely facilitate", company_names)

    def test_english_hospital_coverage_lines_do_not_become_companies(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "Professional Experience",
            "Takeda China",
            "Senior Medical Science Liaison",
            "Jan 2019 - Present",
            "China, covering more than 40 hospital, 40 KOLs, and 230 HCPs.",
            "Therapeutic Area data generation and communication.",
            "Medical Advisor April 2017 - Dec 2018",
            "R.D. community collaboration.",
        ])
        ctx = build_placeholder_context(
            {
                "candidate_name": "Test",
                "position_title": "Medical Science Liaison",
                "original_resume": resume,
                "parsed_resume": parse_resume_for_report(resume),
            },
            {"brand_id": "tstar"},
        )
        company_names = [str(group.get("company") or "") for group in ctx["appendix_blocks"]["experience_groups"]]
        experience_text = " ".join(
            " ".join([str(group["company"]), role.get("period", ""), role.get("title", ""), *role.get("details", [])])
            for group in ctx["appendix_blocks"]["experience_groups"]
            for role in group["roles"]
        )

        self.assertIn("Takeda China", company_names)
        self.assertNotIn("China, covering more than 40 hospital, 40 KOLs, and 230 HCPs.", company_names)
        self.assertNotIn("Therapeutic Area data.", company_names)
        self.assertIn("covering more than 40 hospital", experience_text)

    def test_english_target_hospital_lines_do_not_become_companies(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "Professional Experience",
            "Wyeth China",
            "Product Manager",
            "2012 - 2015",
            "Target hospitals: Pumch, Anzhen Hospital, Puren Hospital.",
            "MKT Strategy Planning and brand life-cycle management.",
            "EFEXOR from Lexapro.",
        ])
        ctx = build_placeholder_context(
            {
                "candidate_name": "Test",
                "position_title": "Product Manager",
                "original_resume": resume,
                "parsed_resume": parse_resume_for_report(resume),
            },
            {"brand_id": "tstar"},
        )
        company_names = [str(group.get("company") or "") for group in ctx["appendix_blocks"]["experience_groups"]]
        experience_text = " ".join(
            " ".join([str(group["company"]), role.get("period", ""), role.get("title", ""), *role.get("details", [])])
            for group in ctx["appendix_blocks"]["experience_groups"]
            for role in group["roles"]
        )

        self.assertIn("Wyeth China", company_names)
        self.assertNotIn("Target hospitals: Pumch, Anzhen Hospital, Puren Hospital.", company_names)
        self.assertIn("Target hospitals: Pumch", experience_text)

    def test_english_be_responsible_hospital_lines_do_not_become_companies(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "Work Experience",
            "Pfizer Investment Co., Ltd.",
            "Senior Product Manager",
            "2017.03-2019.08",
            "Be responsible for Pfizer hospital channel of urology TA, Viagra and Cardura Pfizer.",
            "Hospital and retail strategy and maximize value with 1 Billion RMB.",
            "Xian Janssen Pharmaceutical Ltd.",
            "Medical Specialist",
            "2012.01-2017.02",
        ])
        ctx = build_placeholder_context(
            {
                "candidate_name": "Test",
                "position_title": "Product Manager",
                "original_resume": resume,
                "parsed_resume": parse_resume_for_report(resume),
            },
            {"brand_id": "tstar"},
        )
        company_names = [str(group.get("company") or "") for group in ctx["appendix_blocks"]["experience_groups"]]
        experience_text = " ".join(
            " ".join([str(group["company"]), role.get("period", ""), role.get("title", ""), *role.get("details", [])])
            for group in ctx["appendix_blocks"]["experience_groups"]
            for role in group["roles"]
        )

        self.assertIn("Pfizer Investment Co., Ltd.", company_names)
        self.assertIn("Xian Janssen Pharmaceutical Ltd.", company_names)
        self.assertNotIn("Be responsible for Pfizer hospital channel of urology TA, Viagra and Cardura Pfizer.", company_names)
        self.assertNotIn("Hospital and retail strategy and maximize value with 1 Billion RMB.", company_names)
        self.assertIn("Senior Product Manager", experience_text)
        self.assertIn("Medical Specialist", experience_text)

    def test_english_date_prefixed_company_line_is_cleaned(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "WORKING EXPERIENCE",
            "2019.07-Now InxMed Biotechnology (Shanghai) Co., Ltd",
            "Position: RA Manager",
            "Job Description:",
            "IND submission and management for US, China and Australia.",
            "2016.01-2019.06 Bayer Healthcare Ltd.",
            "Regulatory Affairs Specialist",
        ])
        ctx = build_placeholder_context(
            {
                "candidate_name": "Test",
                "position_title": "RA Manager",
                "original_resume": resume,
                "parsed_resume": parse_resume_for_report(resume),
            },
            {"brand_id": "tstar"},
        )
        group_by_company = {group["company"]: group for group in ctx["appendix_blocks"]["experience_groups"]}

        self.assertIn("InxMed Biotechnology (Shanghai) Co., Ltd", group_by_company)
        self.assertIn("Bayer Healthcare Ltd.", group_by_company)
        self.assertNotIn("2019.07-Now InxMed Biotechnology (Shanghai) Co., Ltd", group_by_company)
        self.assertEqual(group_by_company["InxMed Biotechnology (Shanghai) Co., Ltd"]["roles"][0]["period"], "2019.07-Now")

    def test_english_summary_items_are_not_attached_to_first_company(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "Professional Summary",
            "20+ yrs working with MNCs in medical and life science devices, instrument and consumables.",
            "10+ yrs working as General Manager or national function.",
            "Solid experiences of organization development and distribution management.",
            "Work Experience",
            "Radiometer China",
            "2019.04 - present TCM National Manager, China",
            "Lead nationwide sales and application team members.",
        ])
        ctx = build_placeholder_context(
            {
                "candidate_name": "Test",
                "position_title": "General Manager",
                "original_resume": resume,
                "parsed_resume": parse_resume_for_report(resume),
            },
            {"brand_id": "tstar"},
        )
        experience_text = " ".join(
            " ".join([str(group["company"]), role.get("period", ""), role.get("title", ""), *role.get("details", [])])
            for group in ctx["appendix_blocks"]["experience_groups"]
            for role in group["roles"]
        )

        self.assertIn("Radiometer China", experience_text)
        self.assertNotIn("20+ yrs working with MNCs", experience_text)
        self.assertNotIn("10+ yrs working as General Manager", experience_text)

    def test_english_hospital_company_with_following_period_role_is_separate_experience(self):
        from core.resume_parser import parse_resume_for_report
        from core.placeholder_report import build_placeholder_context

        resume = "\n".join([
            "Work Experience",
            "Olympus Optical Co., Ltd. Shanghai Office",
            "1998.10 - 2000.10 Manager, Surgical Division, East China",
            "In charge of sales and market in surgical endoscopy business.",
            "Shanghai Second People's Hospital",
            "1997.07 - 1998.06 Surgeon",
            "Education",
            "Shanghai Second Medical University",
            "1990.07 - 1995.07 Bachelor of Clinical Medicine",
        ])
        ctx = build_placeholder_context(
            {
                "candidate_name": "Test",
                "position_title": "General Manager",
                "original_resume": resume,
                "parsed_resume": parse_resume_for_report(resume),
            },
            {"brand_id": "tstar"},
        )
        groups = ctx["appendix_blocks"]["experience_groups"]
        group_by_company = {group["company"]: group for group in groups}

        self.assertIn("Olympus Optical Co., Ltd. Shanghai Office", group_by_company)
        self.assertIn("Shanghai Second People's Hospital", group_by_company)
        hospital_role = group_by_company["Shanghai Second People's Hospital"]["roles"][0]
        self.assertEqual(hospital_role["period"], "1997.07 - 1998.06")
        self.assertEqual(hospital_role["title"], "Surgeon")


if __name__ == "__main__":
    unittest.main()
